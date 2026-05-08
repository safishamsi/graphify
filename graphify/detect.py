# file discovery, type classification, and corpus health checks
from __future__ import annotations
import fnmatch
import json
import os
import re
from enum import Enum
from pathlib import Path

from graphify.google_workspace import (
    GOOGLE_WORKSPACE_EXTENSIONS,
    convert_google_workspace_file,
    google_workspace_enabled,
)


class FileType(str, Enum):
    CODE = "code"
    DOCUMENT = "document"
    PAPER = "paper"
    IMAGE = "image"
    VIDEO = "video"


_MANIFEST_PATH = "graphify-out/manifest.json"

CODE_EXTENSIONS = {'.py', '.ts', '.js', '.jsx', '.tsx', '.mjs', '.ejs', '.go', '.rs', '.java', '.groovy', '.gradle', '.cpp', '.cc', '.cxx', '.c', '.h', '.hpp', '.rb', '.swift', '.kt', '.kts', '.cs', '.scala', '.php', '.lua', '.luau', '.toc', '.zig', '.ps1', '.ex', '.exs', '.m', '.mm', '.jl', '.vue', '.svelte', '.dart', '.v', '.sv', '.sql', '.r', '.f', '.F', '.f90', '.F90', '.f95', '.F95', '.f03', '.F03', '.f08', '.F08'}
DOC_EXTENSIONS = {'.md', '.mdx', '.qmd', '.txt', '.rst', '.html', '.yaml', '.yml', '.json', '.jsonc'}
PAPER_EXTENSIONS = {'.pdf'}
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'}
OFFICE_EXTENSIONS = {'.docx', '.xlsx'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.webm', '.mkv', '.avi', '.m4v', '.mp3', '.wav', '.m4a', '.ogg'}

CORPUS_WARN_THRESHOLD = 50_000    # words - below this, warn "you may not need a graph"
CORPUS_UPPER_THRESHOLD = 500_000  # words - above this, warn about token cost
FILE_COUNT_UPPER = 200             # files - above this, warn about token cost

# Files that may contain secrets - skip silently
_SENSITIVE_PATTERNS = [
    re.compile(r'(^|[\\/])\.(env|envrc)(\.|$)', re.IGNORECASE),
    re.compile(r'\.(pem|key|p12|pfx|cert|crt|der|p8)$', re.IGNORECASE),
    re.compile(r'\b(credential|secret|passwd|password|token|private_key)s?\b', re.IGNORECASE),
    re.compile(r'(id_rsa|id_dsa|id_ecdsa|id_ed25519)(\.pub)?$'),
    re.compile(r'(\.netrc|\.pgpass|\.htpasswd)$', re.IGNORECASE),
    re.compile(r'(aws_credentials|gcloud_credentials|service.account)', re.IGNORECASE),
]

# Signals that a .md/.txt file is actually a converted academic paper
_PAPER_SIGNALS = [
    re.compile(r'\barxiv\b', re.IGNORECASE),
    re.compile(r'\bdoi\s*:', re.IGNORECASE),
    re.compile(r'\babstract\b', re.IGNORECASE),
    re.compile(r'\bproceedings\b', re.IGNORECASE),
    re.compile(r'\bjournal\b', re.IGNORECASE),
    re.compile(r'\bpreprint\b', re.IGNORECASE),
    re.compile(r'\\cite\{'),          # LaTeX citation
    re.compile(r'\[\d+\]'),           # Numbered citation [1], [23] (inline)
    re.compile(r'\[\n\d+\n\]'),       # Numbered citation spread across lines (markdown conversion)
    re.compile(r'eq\.\s*\d+|equation\s+\d+', re.IGNORECASE),
    re.compile(r'\d{4}\.\d{4,5}'),   # arXiv ID like 1706.03762
    re.compile(r'\bwe propose\b', re.IGNORECASE),   # common academic phrasing
    re.compile(r'\bliterature\b', re.IGNORECASE),   # "from the literature"
]
_PAPER_SIGNAL_THRESHOLD = 3  # need at least this many signals to call it a paper


def _is_sensitive(path: Path) -> bool:
    """Return True if this file likely contains secrets and should be skipped."""
    name = path.name
    return any(p.search(name) for p in _SENSITIVE_PATTERNS)


def _looks_like_paper(path: Path) -> bool:
    """Heuristic: does this text file read like an academic paper?"""
    try:
        # Only scan first 3000 chars for speed
        text = path.read_text(encoding="utf-8", errors="ignore")[:3000]
        hits = sum(1 for pattern in _PAPER_SIGNALS if pattern.search(text))
        return hits >= _PAPER_SIGNAL_THRESHOLD
    except Exception:
        return False


_ASSET_DIR_MARKERS = {".imageset", ".xcassets", ".appiconset", ".colorset", ".launchimage"}


_SHEBANG_CODE_INTERPRETERS = {
    "python", "python3", "python2",
    "ruby", "perl", "node", "nodejs",
    "bash", "sh", "dash", "zsh", "fish", "ksh", "tcsh",
    "lua", "php", "julia", "Rscript",
}


def _shebang_file_type(path: Path) -> FileType | None:
    """Peek at the first line of an extensionless file for a shebang."""
    try:
        with path.open("rb") as f:
            first = f.read(128)
        if not first.startswith(b"#!"):
            return None
        line = first.split(b"\n")[0].decode(errors="replace")
        parts = line[2:].strip().split()
        if not parts:
            return None
        interp = parts[0].split("/")[-1]  # /usr/bin/env → env
        if interp == "env" and len(parts) > 1:
            interp = parts[1].split("/")[-1]
        if interp in _SHEBANG_CODE_INTERPRETERS:
            return FileType.CODE
    except OSError:
        pass
    return None


def classify_file(path: Path) -> FileType | None:
    # Compound extensions must be checked before simple suffix lookup
    if path.name.lower().endswith(".blade.php"):
        return FileType.CODE
    ext = path.suffix.lower()
    if not ext:
        return _shebang_file_type(path)
    if ext in CODE_EXTENSIONS:
        return FileType.CODE
    if ext in PAPER_EXTENSIONS:
        # PDFs inside Xcode asset catalogs are vector icons, not papers
        if any(part.endswith(tuple(_ASSET_DIR_MARKERS)) for part in path.parts):
            return None
        return FileType.PAPER
    if ext in IMAGE_EXTENSIONS:
        return FileType.IMAGE
    if ext in DOC_EXTENSIONS:
        # Check if it's a converted paper
        if _looks_like_paper(path):
            return FileType.PAPER
        return FileType.DOCUMENT
    if ext in OFFICE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in GOOGLE_WORKSPACE_EXTENSIONS:
        return FileType.DOCUMENT
    if ext in VIDEO_EXTENSIONS:
        return FileType.VIDEO
    return None


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def xlsx_extract_structure(path: Path) -> dict:
    """Extract structural nodes (sheets, named tables, column headers) from an .xlsx file.

    Returns a nodes/edges dict compatible with the graphify extract pipeline.
    Used in addition to xlsx_to_markdown so Claude sees both structure and content.
    """
    def _nid(*parts: str) -> str:
        return re.sub(r"[^a-z0-9_]", "_", "_".join(p.lower() for p in parts).strip("_"))

    try:
        import openpyxl
    except ImportError:
        return {"nodes": [], "edges": []}

    try:
        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
    except Exception:
        return {"nodes": [], "edges": []}

    # F-035: typo fix — was `_re.sub` (NameError, but unreachable because the
    # whole xlsx codepath is currently behind a feature flag / not yet wired
    # into the dispatcher). Before re-enabling this path, re-audit it for
    # zip/XML bombs (openpyxl is built on top of zipfile and lxml-style XML
    # parsing — a malicious .xlsx can blow up memory at load_workbook time).
    stem = re.sub(r"[^a-z0-9]", "_", path.stem.lower())
    str_path = str(path)
    file_nid = _nid(str_path)
    nodes: list[dict] = [{"id": file_nid, "label": path.name, "file_type": "document",
                           "source_file": str_path, "source_location": None}]
    edges: list[dict] = []
    seen: set[str] = {file_nid}

    def _add(nid: str, label: str) -> None:
        if nid not in seen:
            seen.add(nid)
            nodes.append({"id": nid, "label": label, "file_type": "document",
                           "source_file": str_path, "source_location": None})

    def _edge(src: str, tgt: str, relation: str) -> None:
        edges.append({"source": src, "target": tgt, "relation": relation,
                       "confidence": "EXTRACTED", "source_file": str_path,
                       "source_location": None, "weight": 1.0})

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_nid = _nid(stem, sheet_name)
        _add(sheet_nid, f"{sheet_name} (sheet)")
        _edge(file_nid, sheet_nid, "contains")

        # Named Excel Tables (ListObjects)
        if hasattr(ws, "tables"):
            for tbl in ws.tables.values():
                tbl_nid = _nid(stem, sheet_name, tbl.name)
                _add(tbl_nid, tbl.name)
                _edge(sheet_nid, tbl_nid, "contains")
                # Column headers from table header row
                ref = tbl.ref  # e.g. "A1:D10"
                if ref:
                    try:
                        from openpyxl.utils import range_boundaries
                        min_col, min_row, max_col, _ = range_boundaries(ref)
                        header_row = list(ws.iter_rows(min_row=min_row, max_row=min_row,
                                                       min_col=min_col, max_col=max_col,
                                                       values_only=True))
                        if header_row:
                            for col_name in header_row[0]:
                                if col_name:
                                    col_nid = _nid(stem, tbl.name, str(col_name))
                                    _add(col_nid, str(col_name))
                                    _edge(tbl_nid, col_nid, "contains")
                    except Exception:
                        pass
        else:
            # Fallback: first non-empty row as column headers
            for row in ws.iter_rows(max_row=1, values_only=True):
                for cell in row:
                    if cell:
                        col_nid = _nid(stem, sheet_name, str(cell))
                        _add(col_nid, str(cell))
                        _edge(sheet_nid, col_nid, "contains")
                break

    try:
        wb.close()
    except Exception:
        pass

    return {"nodes": nodes, "edges": edges}


def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions
    import hashlib
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path


def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        return len(path.read_text(encoding="utf-8", errors="ignore").split())
    except Exception:
        return 0


# Directory names to always skip - venvs, caches, build artifacts, deps
_SKIP_DIRS = {
    "venv", ".venv", "env", ".env",
    "node_modules", "__pycache__", ".git",
    "dist", "build", "target", "out",
    "site-packages", "lib64",
    ".pytest_cache", ".mypy_cache", ".ruff_cache",
    ".tox", ".eggs", "*.egg-info",
    "graphify-out",  # never treat own output as source input (#524)
}

# Large generated files that are never useful to extract
_SKIP_FILES = {
    "package-lock.json", "yarn.lock", "pnpm-lock.yaml",
    "Cargo.lock", "poetry.lock", "Gemfile.lock",
    "composer.lock", "go.sum", "go.work.sum",
}

# Patterns for generated/bundled files that produce noisy extraction (#728)
_GENERATED_FILE_PATTERNS = (
    "*.min.js",
    "*.min.css",
    "*.bundle.js",
    "*.bundle.css",
    "*.map",
    "*.d.ts",
    "*.generated.*",
    "*.gen.*",
    "*-generated.*",
    "*_generated.*",
)


def _is_generated_or_bundle(path: Path, root: Path, include_patterns: list[tuple[Path, str]]) -> bool:
    """Return True for generated/bundled files that should be skipped by default.

    #728's root cause is not a graph algorithm failure; it is noisy source
    files entering extraction. .graphifyinclude remains the escape hatch for
    teams that intentionally want a generated declaration or bundle indexed.
    """
    if _is_included(path, root, include_patterns):
        return False
    name = path.name.lower()
    try:
        rel = str(path.relative_to(root)).replace(os.sep, "/").lower()
    except ValueError:
        rel = name
    # fnmatch is used intentionally for both the basename and normalized
    # project-relative path. Keep patterns conservative and document
    # .graphifyinclude as the false-positive escape hatch.
    return any(fnmatch.fnmatch(name, pattern) or fnmatch.fnmatch(rel, pattern) for pattern in _GENERATED_FILE_PATTERNS)

def _is_noise_dir(part: str) -> bool:
    """Return True if this directory name looks like a venv, cache, or dep dir."""
    if part in _SKIP_DIRS:
        return True
    # Catch *_venv, *_repo/site-packages patterns
    if part.endswith("_venv") or part.endswith("_env"):
        return True
    if part.endswith(".egg-info"):
        return True
    return False


_VCS_MARKERS = (".git", ".hg", ".svn", "_darcs", ".fossil")


def _parse_gitignore_line(raw: str) -> str:
    """Parse one raw line from a .graphifyignore file per gitignore spec.

    - Strip newline chars
    - Strip inline comments (whitespace + # suffix), but only when # is
      preceded by whitespace — so path#with#hash.py is preserved
    - Unescape \\# to literal #
    - Remove trailing spaces unless escaped with backslash
    - Strip leading whitespace
    - Return empty string for blank lines and full-line comments
    """
    line = raw.rstrip("\n\r")
    line = line.lstrip()
    if not line or line.startswith("#"):
        return ""
    # Strip inline comments: require whitespace before # (gitignore extension)
    line = re.sub(r"\s+#+[^\\].*$", "", line)
    # Unescape \# → literal #
    line = line.replace("\\#", "#")
    # Remove unescaped trailing spaces (per gitignore spec)
    line = re.sub(r"(?<!\\) +$", "", line)
    return line


def _find_vcs_root(start: Path) -> Path | None:
    """Walk upward from start; return the first directory containing a VCS marker."""
    current = start.resolve()
    home = Path.home()
    while True:
        if any((current / m).exists() for m in _VCS_MARKERS):
            return current
        parent = current.parent
        if parent == current or current == home:
            return None
        current = parent


def _load_graphifyignore(root: Path) -> list[tuple[Path, str]]:
    """Read .graphifyignore files and return (anchor_dir, pattern) pairs.

    Patterns are returned outer-first so that inner (closer) rules are
    appended last and win via last-match-wins semantics — matching gitignore
    behavior exactly.

    Walk ceiling: the nearest VCS root if inside a repo, otherwise the scan
    root itself (hermetic — no leakage across unrelated sibling projects).
    """
    root = root.resolve()
    ceiling = _find_vcs_root(root) or root

    # Collect ancestor dirs from ceiling down to root (outer → inner)
    dirs: list[Path] = []
    current = root
    while True:
        dirs.append(current)
        if current == ceiling:
            break
        current = current.parent
    dirs.reverse()  # ceiling first, scan root last

    patterns: list[tuple[Path, str]] = []
    for d in dirs:
        ignore_file = d / ".graphifyignore"
        if ignore_file.exists():
            for raw in ignore_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = _parse_gitignore_line(raw)
                if line:
                    patterns.append((d, line))
    return patterns


def _is_ignored(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    """Return True if the path should be ignored per .graphifyignore patterns.

    Uses gitignore last-match-wins semantics: all patterns are evaluated in
    order; the final matching pattern determines the result. Negation patterns
    (starting with !) un-ignore a previously ignored path.
    """
    if not patterns:
        return False

    def _matches(rel: str, p: str) -> bool:
        parts = rel.split("/")
        if fnmatch.fnmatch(rel, p):
            return True
        if fnmatch.fnmatch(path.name, p):
            return True
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, p):
                return True
            if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                return True
        return False

    result = False
    for anchor, pattern in patterns:
        negated = pattern.startswith("!")
        raw = pattern[1:] if negated else pattern
        anchored = raw.startswith("/")
        p = raw.strip("/")
        if not p:
            continue

        matched = False
        if anchored:
            try:
                rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                matched = _matches(rel_anchor, p)
            except ValueError:
                pass
        else:
            try:
                rel = str(path.relative_to(root)).replace(os.sep, "/")
                matched = _matches(rel, p)
            except ValueError:
                pass
            if not matched and anchor != root:
                try:
                    rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                    matched = _matches(rel_anchor, p)
                except ValueError:
                    pass

        if matched:
            result = not negated  # last match wins; ! flips to un-ignore
    return result


def _load_graphifyinclude(root: Path) -> list[tuple[Path, str]]:
    """Read .graphifyinclude allowlist patterns from root and ancestors.

    Include patterns opt matching hidden files/dirs into traversal. Sensitive
    files and hard-skipped noise directories are still excluded later.
    Uses the same VCS-root ceiling logic as _load_graphifyignore.
    """
    root = root.resolve()
    ceiling = _find_vcs_root(root) or root

    dirs: list[Path] = []
    current = root
    while True:
        dirs.append(current)
        if current == ceiling:
            break
        current = current.parent
    dirs.reverse()

    patterns: list[tuple[Path, str]] = []
    for d in dirs:
        include_file = d / ".graphifyinclude"
        if include_file.exists():
            for raw in include_file.read_text(encoding="utf-8", errors="ignore").splitlines():
                line = _parse_gitignore_line(raw)
                if line:
                    patterns.append((d, line))
    return patterns


def _is_included(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    """Return True if path matches any .graphifyinclude allowlist pattern."""
    if not patterns:
        return False

    def _matches(rel: str, p: str) -> bool:
        parts = rel.split("/")
        if fnmatch.fnmatch(rel, p):
            return True
        if fnmatch.fnmatch(path.name, p):
            return True
        for i, part in enumerate(parts):
            if fnmatch.fnmatch(part, p):
                return True
            if fnmatch.fnmatch("/".join(parts[:i + 1]), p):
                return True
        return False

    for anchor, pattern in patterns:
        anchored = pattern.startswith("/")
        p = pattern.strip("/")
        if not p:
            continue
        if anchored:
            try:
                rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                if _matches(rel_anchor, p):
                    return True
            except ValueError:
                pass
        else:
            try:
                rel = str(path.relative_to(root)).replace(os.sep, "/")
                if _matches(rel, p):
                    return True
            except ValueError:
                pass
            if anchor != root:
                try:
                    rel_anchor = str(path.relative_to(anchor)).replace(os.sep, "/")
                    if _matches(rel_anchor, p):
                        return True
                except ValueError:
                    pass
    return False


def _could_contain_included_path(path: Path, root: Path, patterns: list[tuple[Path, str]]) -> bool:
    """Return True if a directory may contain files matched by .graphifyinclude."""
    if not patterns:
        return False

    rels: list[str] = []
    try:
        rels.append(str(path.relative_to(root)).replace(os.sep, "/"))
    except ValueError:
        pass
    for anchor, _ in patterns:
        if anchor != root:
            try:
                rels.append(str(path.relative_to(anchor)).replace(os.sep, "/"))
            except ValueError:
                pass

    for rel in rels:
        rel = rel.strip("/")
        if not rel:
            return True
        for _, pattern in patterns:
            p = pattern.strip("/")
            if not p:
                continue
            if p == rel or p.startswith(rel + "/"):
                return True
            if fnmatch.fnmatch(rel, p):
                return True
    return False


def detect(root: Path, *, follow_symlinks: bool = False, google_workspace: bool | None = None) -> dict:
    root = root.resolve()
    google_workspace = google_workspace_enabled() if google_workspace is None else google_workspace
    files: dict[FileType, list[str]] = {
        FileType.CODE: [],
        FileType.DOCUMENT: [],
        FileType.PAPER: [],
        FileType.IMAGE: [],
        FileType.VIDEO: [],
    }
    total_words = 0

    skipped_sensitive: list[str] = []
    ignore_patterns = _load_graphifyignore(root)
    include_patterns = _load_graphifyinclude(root)

    # Always include graphify-out/memory/ - query results filed back into the graph
    memory_dir = root / "graphify-out" / "memory"
    scan_paths = [root]
    if memory_dir.exists():
        scan_paths.append(memory_dir)

    seen: set[Path] = set()
    all_files: list[Path] = []

    for scan_root in scan_paths:
        in_memory_tree = memory_dir.exists() and str(scan_root).startswith(str(memory_dir))
        for dirpath, dirnames, filenames in os.walk(scan_root, followlinks=follow_symlinks):
            dp = Path(dirpath)
            if follow_symlinks and os.path.islink(dirpath):
                real = os.path.realpath(dirpath)
                parent_real = os.path.realpath(os.path.dirname(dirpath))
                if parent_real == real or parent_real.startswith(real + os.sep):
                    dirnames.clear()
                    continue
            if not in_memory_tree:
                # Prune noise dirs in-place so os.walk never descends into them.
                # Hidden dirs are allowed through if they could contain an
                # explicitly included path (.graphifyinclude allowlist).
                # When negation patterns (!) exist, skip directory-level ignore
                # pruning so negated files inside can still be reached.
                has_negation = any(p.startswith("!") for _, p in ignore_patterns)
                dirnames[:] = [
                    d for d in dirnames
                    if (not d.startswith(".") or _could_contain_included_path(dp / d, root, include_patterns))
                    and not _is_noise_dir(d)
                    and (has_negation or not _is_ignored(dp / d, root, ignore_patterns))
                ]
            for fname in filenames:
                if fname in _SKIP_FILES:
                    continue
                p = dp / fname
                if p not in seen:
                    seen.add(p)
                    all_files.append(p)

    converted_dir = root / "graphify-out" / "converted"

    for p in all_files:
        # For memory dir files, skip hidden/noise filtering
        in_memory = memory_dir.exists() and str(p).startswith(str(memory_dir))
        if not in_memory:
            # Hidden files are already excluded via dir pruning above,
            # but catch hidden files at the root level. A .graphifyinclude
            # entry can opt a specific hidden file back in.
            if p.name.startswith(".") and not _is_included(p, root, include_patterns):
                continue
            # Skip files inside our own converted/ dir (avoid re-processing sidecars)
            if str(p).startswith(str(converted_dir)):
                continue
            # Skip generated/bundled files unless explicitly included (#728)
            if _is_generated_or_bundle(p, root, include_patterns):
                continue
        if _is_ignored(p, root, ignore_patterns):
            continue
        if _is_sensitive(p):
            skipped_sensitive.append(str(p))
            continue
        ftype = classify_file(p)
        if ftype:
            if p.suffix.lower() in GOOGLE_WORKSPACE_EXTENSIONS:
                if not google_workspace:
                    skipped_sensitive.append(
                        str(p)
                        + " [Google Workspace shortcut skipped - pass --google-workspace "
                        "or set GRAPHIFY_GOOGLE_WORKSPACE=1]"
                    )
                    continue
                try:
                    md_path = convert_google_workspace_file(p, converted_dir, xlsx_to_markdown=xlsx_to_markdown)
                except Exception as exc:
                    skipped_sensitive.append(str(p) + f" [Google Workspace export failed: {exc}]")
                    continue
                if md_path:
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    skipped_sensitive.append(str(p) + " [Google Workspace export produced no readable text]")
                continue
            # Office files: convert to markdown sidecar so subagents can read them
            if p.suffix.lower() in OFFICE_EXTENSIONS:
                md_path = convert_office_file(p, converted_dir)
                if md_path:
                    files[ftype].append(str(md_path))
                    total_words += count_words(md_path)
                else:
                    # Conversion failed (library not installed) - skip with note
                    skipped_sensitive.append(str(p) + " [office conversion failed - pip install graphifyy[office]]")
                continue
            files[ftype].append(str(p))
            if ftype != FileType.VIDEO:
                total_words += count_words(p)

    total_files = sum(len(v) for v in files.values())
    needs_graph = total_words >= CORPUS_WARN_THRESHOLD

    # Determine warning - lower bound, upper bound, or sensitive files skipped
    warning: str | None = None
    if not needs_graph:
        warning = (
            f"Corpus is ~{total_words:,} words - fits in a single context window. "
            f"You may not need a graph."
        )
    elif total_words >= CORPUS_UPPER_THRESHOLD or total_files >= FILE_COUNT_UPPER:
        warning = (
            f"Large corpus: {total_files} files · ~{total_words:,} words. "
            f"Semantic extraction will be expensive (many Claude tokens). "
            f"Consider running on a subfolder, or use --no-semantic to run AST-only."
        )

    return {
        "files": {k.value: v for k, v in files.items()},
        "total_files": total_files,
        "total_words": total_words,
        "needs_graph": needs_graph,
        "warning": warning,
        "skipped_sensitive": skipped_sensitive,
        "graphifyignore_patterns": len(ignore_patterns),
    }


def _md5_file(path: Path) -> str:
    """MD5 of file contents streamed in 64KB chunks — for change detection only."""
    import hashlib as _hl
    h = _hl.md5(usedforsecurity=False)
    try:
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(65536), b""):
                h.update(chunk)
    except OSError:
        return ""
    return h.hexdigest()


def _manifest_root(manifest_path: str | Path) -> Path:
    """Infer the scan root for graphify-out/manifest.json.

    Current code writes raw detect() paths as manifest keys. detect() resolves
    the scan root, so those keys are absolute and break after clone (#777).
    Prefer the sibling .graphify_root value because it records the user's scan
    root. Fall back to the output directory parent for first-run and legacy
    cases where .graphify_root is missing.
    """
    manifest = Path(manifest_path)
    out_dir = manifest.parent
    root_file = out_dir / ".graphify_root"
    if root_file.exists():
        raw_root = root_file.read_text(encoding="utf-8").strip()
        if raw_root:
            saved_root = Path(raw_root)
            return saved_root.resolve() if saved_root.is_absolute() else (out_dir.parent / saved_root).resolve()
    if out_dir.name:
        return out_dir.parent.resolve()
    return Path(".").resolve()


def _manifest_key(path: str | Path, root: Path) -> str:
    """Return the portable manifest key for a source path."""
    root = root.resolve()
    p = Path(path)
    if not p.is_absolute():
        p = (root / p).resolve()
    else:
        p = p.resolve()
    try:
        return str(p.relative_to(root))
    except ValueError:
        # Files outside the scan root are rare but possible via explicit inputs.
        # Keep them absolute rather than manufacturing a misleading relpath.
        return str(p)


def _manifest_path(key: str, root: Path) -> Path:
    """Anchor a manifest key back to an actual path."""
    p = Path(key)
    return p if p.is_absolute() else root.resolve() / p


def _manifest_entry_for(path: str | Path, manifest: dict, root: Path) -> object | None:
    """Find a manifest entry using new relative keys and legacy absolute keys.

    The *manifest* is expected to already be normalized by :func:`load_manifest`
    (which resolves legacy absolute keys).  We still try multiple candidate forms
    for the incoming *path* so that callers are not forced to resolve beforehand.
    """
    p = Path(path)
    root = root.resolve()
    candidates: list[str] = [str(p), _manifest_key(p, root)]
    if not p.is_absolute():
        candidates.append(str((root / p).resolve()))
    else:
        resolved = str(p.resolve())
        if resolved != str(p):
            candidates.append(resolved)
    for key in candidates:
        if key in manifest:
            return manifest[key]
    return None


def load_manifest(manifest_path: str = _MANIFEST_PATH) -> dict:
    """Load the manifest from a previous run. Returns {} on any error.

    Legacy manifests stored absolute path keys (e.g. from ``str(tmp_path/file)``).
    On macOS ``/var`` is a symlink to ``/private/var``, so the same file can
    appear as ``/var/...`` (unresolved) or ``/private/var/...`` (resolved).
    Normalizing keys through ``Path.resolve()`` at load time prevents spurious
    cache misses when the manifest was written with unresolved paths but
    :func:`detect` returns resolved paths (or vice versa).
    """
    try:
        manifest = json.loads(Path(manifest_path).read_text(encoding="utf-8"))
    except Exception:
        return {}
    # Normalize legacy absolute keys so they match the resolved paths coming
    # from detect() (which resolves the scan root).
    out: dict = {}
    for key, value in manifest.items():
        try:
            pk = Path(key)
            if pk.is_absolute():
                out[str(pk.resolve())] = value
            else:
                out[key] = value
        except Exception:
            out[key] = value
    return out


def save_manifest(
    files: dict[str, list[str]],
    manifest_path: str = _MANIFEST_PATH,
    *,
    root: Path | None = None,
) -> None:
    """Save current file mtimes + content hashes for change detection.

    Before #777 this function used the raw file string as the JSON key. Because
    detect() returns absolute paths, committed manifests leaked user paths and
    missed after clone. This version writes root-relative keys when possible.
    """
    root = (root or _manifest_root(manifest_path)).resolve()
    manifest: dict[str, dict] = {}
    for file_list in files.values():
        for f in file_list:
            try:
                p = Path(f)
                if not p.is_absolute():
                    p = root / p
                key = _manifest_key(p, root)
                manifest[key] = {"mtime": p.stat().st_mtime, "hash": _md5_file(p)}
            except OSError:
                pass
    target = Path(manifest_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")


def detect_incremental(
    root: Path,
    manifest_path: str = _MANIFEST_PATH,
    *,
    follow_symlinks: bool = False,
    google_workspace: bool | None = None,
    full: dict | None = None,
) -> dict:
    """Like detect(), but returns only new or modified files since the last run.

    This remains backwards compatible with legacy absolute-key manifests while
    supporting the portable root-relative keys written by save_manifest().
    Callers that already ran detect() can pass ``full`` to avoid a second full
    filesystem scan in the no-op update fast path (#741).
    """
    root = Path(root).resolve()
    if full is None:
        full = detect(root, follow_symlinks=follow_symlinks, google_workspace=google_workspace)
    manifest = load_manifest(manifest_path)

    if not manifest:
        full["incremental"] = True
        full["new_files"] = full["files"]
        full["unchanged_files"] = {k: [] for k in full["files"]}
        full["new_total"] = full["total_files"]
        full["deleted_files"] = []
        return full

    new_files: dict[str, list[str]] = {k: [] for k in full["files"]}
    unchanged_files: dict[str, list[str]] = {k: [] for k in full["files"]}

    for ftype, file_list in full["files"].items():
        for f in file_list:
            stored = _manifest_entry_for(f, manifest, root)
            p = Path(f)
            try:
                current_mtime = p.stat().st_mtime
            except Exception:
                current_mtime = 0

            if isinstance(stored, (int, float)):
                changed = stored is None or current_mtime > stored
            elif isinstance(stored, dict):
                stored_mtime = stored.get("mtime")
                stored_hash = stored.get("hash", "")
                if stored_mtime is None or current_mtime != stored_mtime:
                    changed = not stored_hash or _md5_file(p) != stored_hash
                else:
                    # mtime unchanged — still verify hash to catch rapid writes
                    # or tools that preserve mtimes (#F7)
                    changed = bool(stored_hash) and _md5_file(p) != stored_hash
            else:
                changed = True

            if changed:
                new_files[ftype].append(f)
            else:
                unchanged_files[ftype].append(f)

    current_keys = {
        _manifest_key(f, root)
        for flist in full["files"].values()
        for f in flist
    }
    current_legacy = {
        str(Path(f))
        for flist in full["files"].values()
        for f in flist
    }
    deleted_files = [
        str(_manifest_path(key, root))
        for key in manifest
        if key not in current_keys
        and key not in current_legacy
        and not _manifest_path(key, root).exists()
    ]

    new_total = sum(len(v) for v in new_files.values())
    full["incremental"] = True
    full["new_files"] = new_files
    full["unchanged_files"] = unchanged_files
    full["new_total"] = new_total
    full["deleted_files"] = deleted_files
    return full
