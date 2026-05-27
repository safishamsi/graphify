import re
from pathlib import Path

def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""

def docx_to_markdown(path: Path) -> str:
    """Convert a .docx file to markdown text using python-docx."""
    try:
        from docx import Document
        doc = Document(str(path))
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            if style.startswith("Heading 1"):
                lines.append(f"# {text}")
            elif style.startswith("Heading 2"):
                lines.append(f"## {text}")
            elif style.startswith("Heading 3"):
                lines.append(f"### {text}")
            elif style.startswith("List"):
                lines.append(f"- {text}")
            else:
                lines.append(text)
        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
    except ImportError:
        return ""
    except Exception:
        return ""

def xlsx_to_markdown(path: Path) -> str:
    """Convert an .xlsx file to markdown text using openpyxl."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue
            sections.append(f"## Sheet: {sheet_name}")
            if len(rows) >= 1:
                header = "| " + " | ".join(rows[0]) + " |"
                sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
                sections.extend([header, sep])
                for row in rows[1:]:
                    sections.append("| " + " | ".join(row) + " |")
        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""

def xlsx_extract_structure(path: Path) -> dict:
    """Extract structural nodes (sheets, named tables, column headers) from an .xlsx file.

    Returns a nodes/edges dict compatible with the graphify extract pipeline.
    Used in addition to xlsx_to_markdown so Claude sees both structure and content.
    """
    def _nid(*parts: str) -> str:
        return re.sub(r"[^a-z0-9_]", "_", "_".join(p.lower() for p in parts).strip("_"))

    try:
        import openpyxl
    except ImportError:
        return {"nodes": [], "edges": []}

    try:
        wb = openpyxl.load_workbook(str(path), read_only=False, data_only=True)
    except Exception:
        return {"nodes": [], "edges": []}

    # F-035: typo fix — was `_re.sub` (NameError, but unreachable because the
    # whole xlsx codepath is currently behind a feature flag / not yet wired
    # into the dispatcher). Before re-enabling this path, re-audit it for
    # zip/XML bombs (openpyxl is built on top of zipfile and lxml-style XML
    # parsing — a malicious .xlsx can blow up memory at load_workbook time).
    stem = re.sub(r"[^a-z0-9]", "_", path.stem.lower())
    str_path = str(path)
    file_nid = _nid(str_path)
    nodes: list[dict] = [{"id": file_nid, "label": path.name, "file_type": "document",
                           "source_file": str_path, "source_location": None}]
    edges: list[dict] = []
    seen: set[str] = {file_nid}

    def _add(nid: str, label: str) -> None:
        if nid not in seen:
            seen.add(nid)
            nodes.append({"id": nid, "label": label, "file_type": "document",
                           "source_file": str_path, "source_location": None})

    def _edge(src: str, tgt: str, relation: str) -> None:
        edges.append({"source": src, "target": tgt, "relation": relation,
                       "confidence": "EXTRACTED", "source_file": str_path,
                       "source_location": None, "weight": 1.0})

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_nid = _nid(stem, sheet_name)
        _add(sheet_nid, f"{sheet_name} (sheet)")
        _edge(file_nid, sheet_nid, "contains")

        # Named Excel Tables (ListObjects)
        if hasattr(ws, "tables"):
            for tbl in ws.tables.values():
                tbl_nid = _nid(stem, sheet_name, tbl.name)
                _add(tbl_nid, tbl.name)
                _edge(sheet_nid, tbl_nid, "contains")
                # Column headers from table header row
                ref = tbl.ref  # e.g. "A1:D10"
                if ref:
                    try:
                        from openpyxl.utils import range_boundaries
                        min_col, min_row, max_col, _ = range_boundaries(ref)
                        header_row = list(ws.iter_rows(min_row=min_row, max_row=min_row,
                                                       min_col=min_col, max_col=max_col,
                                                       values_only=True))
                        if header_row:
                            for col_name in header_row[0]:
                                if col_name:
                                    col_nid = _nid(stem, tbl.name, str(col_name))
                                    _add(col_nid, str(col_name))
                                    _edge(tbl_nid, col_nid, "contains")
                    except Exception:
                        pass
        else:
            # Fallback: first non-empty row as column headers
            for row in ws.iter_rows(max_row=1, values_only=True):
                for cell in row:
                    if cell:
                        col_nid = _nid(stem, sheet_name, str(cell))
                        _add(col_nid, str(cell))
                        _edge(sheet_nid, col_nid, "contains")
                break

    try:
        wb.close()
    except Exception:
        pass

    return {"nodes": nodes, "edges": edges}

def convert_office_file(path: Path, out_dir: Path) -> Path | None:
    """Convert a .docx or .xlsx to a markdown sidecar in out_dir.

    Returns the path of the converted .md file, or None if conversion failed
    or the required library is not installed.
    """
    ext = path.suffix.lower()
    if ext == ".docx":
        text = docx_to_markdown(path)
    elif ext == ".xlsx":
        text = xlsx_to_markdown(path)
    else:
        return None

    if not text.strip():
        return None

    out_dir.mkdir(parents=True, exist_ok=True)
    # Use a stable name derived from the original path to avoid collisions
    import hashlib
    name_hash = hashlib.sha256(str(path.resolve()).encode()).hexdigest()[:8]
    out_path = out_dir / f"{path.stem}_{name_hash}.md"
    out_path.write_text(
        f"<!-- converted from {path.name} -->\n\n{text}",
        encoding="utf-8",
    )
    return out_path

def count_words(path: Path) -> int:
    try:
        ext = path.suffix.lower()
        if ext == ".pdf":
            return len(extract_pdf_text(path).split())
        if ext == ".docx":
            return len(docx_to_markdown(path).split())
        if ext == ".xlsx":
            return len(xlsx_to_markdown(path).split())
        return len(path.read_text(encoding="utf-8", errors="ignore").split())
    except Exception:
        return 0
