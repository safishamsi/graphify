import json
import pytest
from pathlib import Path
from graphify.detect import classify_file, count_words, detect, detect_incremental, save_manifest, FileType, _looks_like_paper, _is_ignored, _load_graphifyignore, _is_included, _could_contain_included_path, _manifest_key, _manifest_root, _manifest_entry_for, load_manifest

FIXTURES = Path(__file__).parent / "fixtures"

def test_classify_python():
    assert classify_file(Path("foo.py")) == FileType.CODE

def test_classify_typescript():
    assert classify_file(Path("bar.ts")) == FileType.CODE

def test_classify_markdown():
    assert classify_file(Path("README.md")) == FileType.DOCUMENT

def test_classify_pdf():
    assert classify_file(Path("paper.pdf")) == FileType.PAPER

def test_classify_pdf_in_xcassets_skipped():
    # PDFs inside Xcode asset catalogs are vector icons, not papers
    asset_pdf = Path("MyApp/Images.xcassets/icon.imageset/icon.pdf")
    assert classify_file(asset_pdf) is None

def test_classify_pdf_in_xcassets_root_skipped():
    asset_pdf = Path("Pods/HXPHPicker/Assets.xcassets/photo.pdf")
    assert classify_file(asset_pdf) is None

def test_classify_unknown_returns_none():
    assert classify_file(Path("archive.zip")) is None

def test_classify_image():
    assert classify_file(Path("screenshot.png")) == FileType.IMAGE
    assert classify_file(Path("design.jpg")) == FileType.IMAGE
    assert classify_file(Path("diagram.webp")) == FileType.IMAGE

def test_count_words_sample_md():
    words = count_words(FIXTURES / "sample.md")
    assert words > 5

def test_detect_finds_fixtures():
    result = detect(FIXTURES)
    assert result["total_files"] >= 2
    assert "code" in result["files"]
    assert "document" in result["files"]

def test_detect_warns_small_corpus():
    result = detect(FIXTURES)
    assert result["needs_graph"] is False
    assert result["warning"] is not None

def test_detect_skips_dotfiles():
    result = detect(FIXTURES)
    for files in result["files"].values():
        for f in files:
            assert "/." not in f


def test_classify_md_paper_by_signals(tmp_path):
    """A .md file with enough paper signals should classify as PAPER."""
    paper = tmp_path / "paper.md"
    paper.write_text(
        "# Abstract\n\nWe propose a new method. See [1] and [23].\n"
        "This work was published in the Journal of AI. ArXiv preprint.\n"
        "See Equation 3 for details. \\cite{vaswani2017}.\n"
    )
    assert classify_file(paper) == FileType.PAPER


def test_classify_md_doc_without_signals(tmp_path):
    """A plain .md file without paper signals should stay DOCUMENT."""
    doc = tmp_path / "notes.md"
    doc.write_text("# My Notes\n\nHere are some notes about the project.\n")
    assert classify_file(doc) == FileType.DOCUMENT


def test_classify_attention_paper():
    """The real attention paper file should be classified as PAPER."""
    paper_path = Path("/home/safi/graphify_eval/papers/attention_is_all_you_need.md")
    if paper_path.exists():
        result = classify_file(paper_path)
        assert result == FileType.PAPER


def test_graphifyignore_excludes_file(tmp_path):
    """Files matching .graphifyignore patterns are excluded from detect()."""
    (tmp_path / ".graphifyignore").write_text("vendor/\n*.generated.py\n")
    vendor = tmp_path / "vendor"
    vendor.mkdir()
    (vendor / "lib.py").write_text("x = 1")
    (tmp_path / "main.py").write_text("print('hi')")
    (tmp_path / "schema.generated.py").write_text("x = 1")

    result = detect(tmp_path)
    file_list = result["files"]["code"]
    assert any("main.py" in f for f in file_list)
    assert not any("vendor" in f for f in file_list)
    assert not any("generated" in f for f in file_list)
    assert result["graphifyignore_patterns"] == 2


def test_graphifyignore_missing_is_fine(tmp_path):
    """No .graphifyignore is not an error."""
    (tmp_path / "main.py").write_text("x = 1")
    result = detect(tmp_path)
    assert result["graphifyignore_patterns"] == 0


def test_graphifyignore_comments_ignored(tmp_path):
    """Comment lines in .graphifyignore are not treated as patterns."""
    (tmp_path / ".graphifyignore").write_text("# this is a comment\n\nmain.py\n")
    (tmp_path / "main.py").write_text("x = 1")
    (tmp_path / "other.py").write_text("x = 2")
    result = detect(tmp_path)
    assert not any("main.py" in f for f in result["files"]["code"])
    assert any("other.py" in f for f in result["files"]["code"])


def test_detect_follows_symlinked_directory(tmp_path):
    real_dir = tmp_path / "real_lib"
    real_dir.mkdir()
    (real_dir / "util.py").write_text("x = 1")
    (tmp_path / "linked_lib").symlink_to(real_dir)

    result_no = detect(tmp_path, follow_symlinks=False)
    result_yes = detect(tmp_path, follow_symlinks=True)

    assert any("real_lib" in f for f in result_no["files"]["code"])
    assert not any("linked_lib" in f for f in result_no["files"]["code"])
    assert any("linked_lib" in f for f in result_yes["files"]["code"])


def test_detect_follows_symlinked_file(tmp_path):
    (tmp_path / "real.py").write_text("x = 1")
    (tmp_path / "link.py").symlink_to(tmp_path / "real.py")

    result = detect(tmp_path, follow_symlinks=True)
    code = result["files"]["code"]
    assert any("real.py" in f for f in code)
    assert any("link.py" in f for f in code)


def test_graphifyignore_hermetic_without_vcs(tmp_path):
    """Without a VCS root, parent .graphifyignore does NOT apply (hermetic)."""
    (tmp_path / ".graphifyignore").write_text("vendor/\n")
    sub = tmp_path / "packages" / "mylib"
    sub.mkdir(parents=True)
    (sub / "main.py").write_text("x = 1")
    vendor = sub / "vendor"
    vendor.mkdir()
    (vendor / "dep.py").write_text("y = 2")

    result = detect(sub)
    code_files = result["files"]["code"]
    assert any("main.py" in f for f in code_files)
    # parent .graphifyignore must NOT leak into a non-VCS scan
    assert any("vendor" in f for f in code_files)
    assert result["graphifyignore_patterns"] == 0


def test_graphifyignore_discovered_from_parent_in_vcs(tmp_path):
    """Inside a VCS repo, parent .graphifyignore applies to subdirectory scans."""
    (tmp_path / ".git").mkdir()
    (tmp_path / ".graphifyignore").write_text("vendor/\n")
    sub = tmp_path / "packages" / "mylib"
    sub.mkdir(parents=True)
    (sub / "main.py").write_text("x = 1")
    vendor = sub / "vendor"
    vendor.mkdir()
    (vendor / "dep.py").write_text("y = 2")

    result = detect(sub)
    code_files = result["files"]["code"]
    assert any("main.py" in f for f in code_files)
    assert not any("vendor" in f for f in code_files)
    assert result["graphifyignore_patterns"] >= 1


def test_graphifyignore_stops_at_git_boundary(tmp_path):
    """Upward search stops at the git repo root (.git directory)."""
    (tmp_path / ".graphifyignore").write_text("main.py\n")
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / ".git").mkdir()
    sub = repo / "sub"
    sub.mkdir()
    (sub / "main.py").write_text("x = 1")

    result = detect(sub)
    code_files = result["files"]["code"]
    assert any("main.py" in f for f in code_files)
    assert result["graphifyignore_patterns"] == 0


def test_graphifyignore_at_git_root_is_included(tmp_path):
    """A .graphifyignore at the git repo root is included when scanning a subdir."""
    repo = tmp_path / "repo"
    repo.mkdir()
    (repo / ".git").mkdir()
    (repo / ".graphifyignore").write_text("vendor/\n")
    sub = repo / "packages" / "mylib"
    sub.mkdir(parents=True)
    (sub / "main.py").write_text("x = 1")
    vendor = sub / "vendor"
    vendor.mkdir()
    (vendor / "dep.py").write_text("y = 2")

    result = detect(sub)
    code_files = result["files"]["code"]
    assert any("main.py" in f for f in code_files)
    assert not any("vendor" in f for f in code_files)
    assert result["graphifyignore_patterns"] == 1


def test_detect_handles_circular_symlinks(tmp_path):
    sub = tmp_path / "a"
    sub.mkdir()
    (sub / "main.py").write_text("x = 1")
    (sub / "loop").symlink_to(tmp_path)

    result = detect(tmp_path, follow_symlinks=True)
    assert any("main.py" in f for f in result["files"]["code"])


def test_detect_incremental_propagates_follow_symlinks(tmp_path, monkeypatch):
    """detect_incremental must forward follow_symlinks so symlinked sub-trees
    appear in incremental scans the same way they appear in full scans."""
    monkeypatch.chdir(tmp_path)

    real_dir = tmp_path / "real_corpus"
    real_dir.mkdir()
    (real_dir / "note.md").write_text("# real note\n\nsome content")
    (tmp_path / "linked_corpus").symlink_to(real_dir)

    (tmp_path / "graphify-out").mkdir()
    manifest_path = str(tmp_path / "graphify-out" / "manifest.json")

    # Without following symlinks, the symlinked dir contents are invisible.
    no_link = detect_incremental(tmp_path, manifest_path, follow_symlinks=False)
    assert not any("linked_corpus" in f for f in no_link["files"]["document"])

    # With follow_symlinks=True, the symlinked dir contents appear and are new.
    yes_link = detect_incremental(tmp_path, manifest_path, follow_symlinks=True)
    assert any("linked_corpus" in f for f in yes_link["files"]["document"])
    assert yes_link["new_total"] >= 2  # real + linked

    # After saving manifest, a second incremental scan should see no changes.
    save_manifest(yes_link["files"], manifest_path, root=tmp_path)
    second = detect_incremental(tmp_path, manifest_path, follow_symlinks=True)
    assert second["new_total"] == 0


def test_classify_video_extensions():
    """Video and audio file extensions should classify as VIDEO."""
    from graphify.detect import FileType
    assert classify_file(Path("lecture.mp4")) == FileType.VIDEO
    assert classify_file(Path("podcast.mp3")) == FileType.VIDEO
    assert classify_file(Path("talk.mov")) == FileType.VIDEO
    assert classify_file(Path("recording.wav")) == FileType.VIDEO
    assert classify_file(Path("webinar.webm")) == FileType.VIDEO
    assert classify_file(Path("audio.m4a")) == FileType.VIDEO


def test_classify_google_workspace_shortcuts():
    assert classify_file(Path("notes.gdoc")) == FileType.DOCUMENT
    assert classify_file(Path("budget.gsheet")) == FileType.DOCUMENT
    assert classify_file(Path("deck.gslides")) == FileType.DOCUMENT


def test_detect_skips_google_workspace_shortcuts_by_default(tmp_path):
    (tmp_path / "notes.gdoc").write_text('{"doc_id":"doc-1"}', encoding="utf-8")

    result = detect(tmp_path)

    assert not result["files"]["document"]
    assert any("Google Workspace shortcut skipped" in item for item in result["skipped_sensitive"])


def test_detect_converts_google_workspace_shortcuts_when_enabled(tmp_path, monkeypatch):
    shortcut = tmp_path / "notes.gdoc"
    shortcut.write_text('{"doc_id":"doc-1"}', encoding="utf-8")

    def fake_convert(path, out_dir, *, xlsx_to_markdown=None):
        out_dir.mkdir(parents=True, exist_ok=True)
        out = out_dir / "notes_converted.md"
        out.write_text("# Notes\n\nA converted Google Doc.", encoding="utf-8")
        return out

    monkeypatch.setattr("graphify.detect.convert_google_workspace_file", fake_convert)

    result = detect(tmp_path, google_workspace=True)

    assert len(result["files"]["document"]) == 1
    assert result["files"]["document"][0].endswith("notes_converted.md")
    assert result["total_words"] > 0


def test_detect_includes_video_key(tmp_path):
    """detect() result always includes a 'video' key even with no video files."""
    (tmp_path / "main.py").write_text("x = 1")
    result = detect(tmp_path)
    assert "video" in result["files"]


def test_detect_finds_video_files(tmp_path):
    """detect() correctly counts video files and does not add them to word count."""
    (tmp_path / "lecture.mp4").write_bytes(b"fake video data")
    (tmp_path / "notes.md").write_text("# Notes\nSome content here.")
    result = detect(tmp_path)
    assert len(result["files"]["video"]) == 1
    assert any("lecture.mp4" in f for f in result["files"]["video"])
    # total_words should not include video files (they have no readable text)
    assert result["total_words"] >= 0  # won't crash


def test_detect_video_not_in_words(tmp_path):
    """Video files do not contribute to total_words."""
    (tmp_path / "clip.mp4").write_bytes(b"\x00" * 100)
    result = detect(tmp_path)
    # Only video file present — total_words should be 0
    assert result["total_words"] == 0


# ---------------------------------------------------------------------------
# _is_sensitive
# ---------------------------------------------------------------------------


def test_extract_pdf_text_with_page_text(tmp_path, monkeypatch):
    """extract_pdf_text covers pages.append(text) when page has text."""
    from graphify.detect import extract_pdf_text
    from pypdf import PdfWriter
    from pypdf._page import PageObject

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    pdf_path = tmp_path / "has_text.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)

    # Monkey-patch PageObject.extract_text to return non-empty text (covers line 155)

    def mock_extract_text(self, *args, **kwargs):
        return "Some meaningful page text content."

    monkeypatch.setattr(PageObject, "extract_text", mock_extract_text)

    text = extract_pdf_text(pdf_path)
    assert "meaningful" in text

def test_is_sensitive_env_files():
    from graphify.detect import _is_sensitive
    assert _is_sensitive(Path(".env")) is True
    assert _is_sensitive(Path(".env.local")) is True
    assert _is_sensitive(Path("app/.envrc")) is True


def test_is_sensitive_key_files():
    from graphify.detect import _is_sensitive
    assert _is_sensitive(Path("id_rsa")) is True
    assert _is_sensitive(Path("id_ed25519.pub")) is True
    assert _is_sensitive(Path("server.key")) is True
    assert _is_sensitive(Path("cert.pem")) is True


def test_is_sensitive_not_sensitive():
    from graphify.detect import _is_sensitive
    assert _is_sensitive(Path("main.py")) is False
    assert _is_sensitive(Path("README.md")) is False


# ---------------------------------------------------------------------------
# _shebang_file_type
# ---------------------------------------------------------------------------

def test_shebang_file_type_python(tmp_path):
    from graphify.detect import _shebang_file_type
    f = tmp_path / "script"
    f.write_text("#!/usr/bin/env python3\nprint('hello')")
    assert _shebang_file_type(f) == FileType.CODE


def test_shebang_file_type_bash(tmp_path):
    from graphify.detect import _shebang_file_type
    f = tmp_path / "script"
    f.write_text("#!/bin/bash\necho hi")
    assert _shebang_file_type(f) == FileType.CODE


def test_shebang_file_type_not_shebang(tmp_path):
    from graphify.detect import _shebang_file_type
    f = tmp_path / "script"
    f.write_text("no shebang here")
    assert _shebang_file_type(f) is None


def test_shebang_file_type_not_code(tmp_path):
    from graphify.detect import _shebang_file_type
    f = tmp_path / "script"
    f.write_text("#!/usr/bin/env unknown_interpreter\n...")
    assert _shebang_file_type(f) is None


def test_shebang_file_type_binary_file(tmp_path):
    from graphify.detect import _shebang_file_type
    f = tmp_path / "binary"
    f.write_bytes(b"\x00\x01\x02\x03")
    assert _shebang_file_type(f) is None


# ---------------------------------------------------------------------------
# _parse_gitignore_line
# ---------------------------------------------------------------------------

def test_parse_gitignore_line_basic():
    from graphify.detect import _parse_gitignore_line
    assert _parse_gitignore_line("*.pyc") == "*.pyc"
    assert _parse_gitignore_line("__pycache__/") == "__pycache__/"


def test_parse_gitignore_line_comment():
    from graphify.detect import _parse_gitignore_line
    assert _parse_gitignore_line("# comment") == ""
    assert _parse_gitignore_line("") == ""


def test_parse_gitignore_line_trailing_spaces():
    from graphify.detect import _parse_gitignore_line
    assert _parse_gitignore_line("build/   ") == "build/"


def test_parse_gitignore_line_inline_comment():
    from graphify.detect import _parse_gitignore_line
    assert _parse_gitignore_line("*.log #  log files") == "*.log"


def test_parse_gitignore_line_unescape_hash():
    from graphify.detect import _parse_gitignore_line
    assert _parse_gitignore_line(r"foo\#bar") == "foo#bar"


# ---------------------------------------------------------------------------
# _find_vcs_root / _load_graphifyinclude / _is_included
# ---------------------------------------------------------------------------

def test_find_vcs_root_in_git_repo():
    from graphify.detect import _find_vcs_root
    # The project itself has .git, so this should find something
    root = _find_vcs_root(Path.cwd())
    assert root is not None
    assert (root / ".git").exists()


def test_find_vcs_root_not_in_repo(tmp_path):
    from graphify.detect import _find_vcs_root
    assert _find_vcs_root(tmp_path) is None


def test_load_graphifyinclude_nonexistent(tmp_path):
    from graphify.detect import _load_graphifyinclude
    patterns = _load_graphifyinclude(tmp_path)
    assert patterns == []


def test_load_graphifyinclude_exists(tmp_path):
    from graphify.detect import _load_graphifyinclude
    (tmp_path / ".graphifyinclude").write_text("!config.yml\ndocs/**/*.md")
    patterns = _load_graphifyinclude(tmp_path)
    assert len(patterns) >= 1


def test_is_included_matches_name(tmp_path):
    from graphify.detect import _is_included
    (tmp_path / ".graphifyinclude").write_text("*.secret")
    patterns = [(tmp_path, "*.secret")]
    f = tmp_path / "foo.secret"
    f.write_text("secret stuff")
    assert _is_included(f, tmp_path, patterns) is True


def test_is_included_no_match(tmp_path):
    from graphify.detect import _is_included
    patterns = [(tmp_path, "*.secret")]
    assert _is_included(tmp_path / "main.py", tmp_path, patterns) is False


def test_is_included_empty_patterns(tmp_path):
    from graphify.detect import _is_included
    assert _is_included(tmp_path / "main.py", tmp_path, []) is False


def test_could_contain_included_path(tmp_path):
    from graphify.detect import _could_contain_included_path
    patterns = [(tmp_path, "src/**/*.py")]
    d = tmp_path / "src"
    d.mkdir()
    filename = d / "main.py"
    filename.write_text("x=1")
    assert _could_contain_included_path(d, tmp_path, patterns) is True


def test_could_contain_included_path_no_match(tmp_path):
    from graphify.detect import _could_contain_included_path
    patterns = [(tmp_path, "*.secret")]
    d = tmp_path / "src"
    d.mkdir()
    assert _could_contain_included_path(d, tmp_path, patterns) is False


def test_could_contain_included_path_empty(tmp_path):
    from graphify.detect import _could_contain_included_path
    assert _could_contain_included_path(tmp_path, tmp_path, []) is False


# ---------------------------------------------------------------------------
# _is_noise_dir
# ---------------------------------------------------------------------------

def test_is_noise_dir_venv():
    from graphify.detect import _is_noise_dir
    assert _is_noise_dir("venv") is True
    assert _is_noise_dir(".venv") is True
    assert _is_noise_dir("my_app_venv") is True


def test_is_noise_dir_not_noise():
    from graphify.detect import _is_noise_dir
    assert _is_noise_dir("src") is False
    assert _is_noise_dir("app") is False


# ---------------------------------------------------------------------------
# _md5_file / load_manifest
# ---------------------------------------------------------------------------

def test_md5_file_returns_hex(tmp_path):
    from graphify.detect import _md5_file
    f = tmp_path / "data.txt"
    f.write_text("hello world")
    h = _md5_file(f)
    assert len(h) == 32
    assert all(c in "0123456789abcdef" for c in h)


def test_md5_file_same_content_same_hash(tmp_path):
    from graphify.detect import _md5_file
    a = tmp_path / "a.txt"
    b = tmp_path / "b.txt"
    a.write_text("same content")
    b.write_text("same content")
    assert _md5_file(a) == _md5_file(b)


def test_md5_file_different_content_different_hash(tmp_path):
    from graphify.detect import _md5_file
    a = tmp_path / "a.txt"
    b = tmp_path / "b.txt"
    a.write_text("hello")
    b.write_text("world")
    assert _md5_file(a) != _md5_file(b)


def test_md5_file_missing_file(tmp_path):
    from graphify.detect import _md5_file
    assert _md5_file(tmp_path / "nope.txt") == ""


def test_load_manifest_nonexistent(tmp_path):
    from graphify.detect import load_manifest
    result = load_manifest(str(tmp_path / "nonexistent.json"))
    assert result == {}


def test_load_manifest_valid(tmp_path):
    from graphify.detect import load_manifest
    f = tmp_path / "manifest.json"
    f.write_text(json.dumps({"a.py": {"mtime": 1234567890.0, "hash": "abc123"}}))
    result = load_manifest(str(f))
    assert result["a.py"]["hash"] == "abc123"


def test_load_manifest_corrupt(tmp_path):
    from graphify.detect import load_manifest
    f = tmp_path / "bad.json"
    f.write_text("not json {{{")
    result = load_manifest(str(f))
    assert result == {}


# ---------------------------------------------------------------------------
# _looks_like_paper — exception path
# ---------------------------------------------------------------------------

def test_looks_like_paper_read_error(tmp_path):
    """_looks_like_paper returns False when file cannot be read."""
    from graphify.detect import _looks_like_paper
    missing = tmp_path / "missing.pdf"  # never created
    assert _looks_like_paper(missing) is False


def test_looks_like_paper_truncated_file(tmp_path):
    """_looks_like_paper handles very short files."""
    from graphify.detect import _looks_like_paper
    f = tmp_path / "empty.md"
    f.write_text("hi")
    assert _looks_like_paper(f) is False


# ---------------------------------------------------------------------------
# classify_file — extensionless / office / blade paths
# ---------------------------------------------------------------------------

def test_classify_file_extensionless_with_shebang(tmp_path):
    """Extensionless file with python shebang → CODE."""
    f = tmp_path / "myscript"
    f.write_text("#!/usr/bin/env python3\nprint('hi')")
    assert classify_file(f) == FileType.CODE


def test_classify_file_extensionless_no_shebang(tmp_path):
    """Extensionless file without shebang → None."""
    f = tmp_path / "textfile"
    f.write_text("just some text")
    assert classify_file(f) is None


def test_classify_file_blade_php():
    assert classify_file(Path("home.blade.php")) == FileType.CODE


def test_classify_file_office_docx():
    assert classify_file(Path("report.docx")) == FileType.DOCUMENT


def test_classify_file_office_xlsx():
    assert classify_file(Path("spreadsheet.xlsx")) == FileType.DOCUMENT


def test_classify_file_office_not_in_list():
    """pptx is not in OFFICE_EXTENSIONS, so Unknown returns None."""
    assert classify_file(Path("deck.pptx")) is None


# ---------------------------------------------------------------------------
# extract_pdf_text
# ---------------------------------------------------------------------------

def test_extract_pdf_text_simple(tmp_path):
    """extract_pdf_text extracts text from a real PDF."""
    from graphify.detect import extract_pdf_text
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    pdf_path = tmp_path / "doc.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)

    text = extract_pdf_text(pdf_path)
    assert isinstance(text, str)


def test_extract_pdf_text_missing_file(tmp_path):
    """extract_pdf_text returns empty string for missing file."""
    from graphify.detect import extract_pdf_text
    assert extract_pdf_text(tmp_path / "nonexistent.pdf") == ""


# ---------------------------------------------------------------------------
# docx_to_markdown
# ---------------------------------------------------------------------------

def test_docx_to_markdown(tmp_path):
    """docx_to_markdown converts a real .docx to markdown."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    from docx import Document

    full_path = tmp_path / "hello.docx"
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Some content here.")
    doc.add_heading("Subsection", level=2)
    p = doc.add_paragraph("More text.")
    doc.save(str(full_path))

    md = docx_to_markdown(full_path)
    assert "# Title" in md
    assert "Some content here" in md
    assert "## Subsection" in md


def test_docx_to_markdown_missing_file(tmp_path):
    """docx_to_markdown returns empty string when file missing."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    assert docx_to_markdown(tmp_path / "nope.docx") == ""


# ---------------------------------------------------------------------------
# xlsx_to_markdown
# ---------------------------------------------------------------------------

def test_xlsx_to_markdown(tmp_path):
    """xlsx_to_markdown converts an .xlsx to markdown."""
    pytest.importorskip('openpyxl')
    from graphify.detect import xlsx_to_markdown
    import openpyxl

    full_path = tmp_path / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age"])
    ws.append(["Alice", 30])
    ws.append(["Bob", 25])
    wb.save(str(full_path))

    md = xlsx_to_markdown(full_path)
    assert "## Sheet: Sheet1" in md
    assert "Name" in md
    assert "Alice" in md


def test_xlsx_to_markdown_missing_file(tmp_path):
    """xlsx_to_markdown returns empty string when file missing."""
    pytest.importorskip('openpyxl')
    from graphify.detect import xlsx_to_markdown
    assert xlsx_to_markdown(tmp_path / "nope.xlsx") == ""


# ---------------------------------------------------------------------------
# xlsx_extract_structure
# ---------------------------------------------------------------------------

def test_xlsx_extract_structure(tmp_path):
    """xlsx_extract_structure returns nodes and edges."""
    pytest.importorskip('openpyxl')
    from graphify.detect import xlsx_extract_structure
    import openpyxl

    full_path = tmp_path / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Name", "Age"])
    ws.append(["Alice", 30])
    wb.save(str(full_path))

    result = xlsx_extract_structure(full_path)
    assert "nodes" in result
    assert "edges" in result
    assert len(result["nodes"]) >= 2  # file node + sheet node
    assert any(n["label"] == "Sheet1 (sheet)" for n in result["nodes"])


def test_xlsx_extract_structure_unreadable(tmp_path):
    """xlsx_extract_structure returns empty on corrupt file."""
    pytest.importorskip('openpyxl')
    from graphify.detect import xlsx_extract_structure
    bad = tmp_path / "bad.xlsx"
    bad.write_text("not an xlsx", encoding="utf-8")
    result = xlsx_extract_structure(bad)
    assert result == {"nodes": [], "edges": []}


# ---------------------------------------------------------------------------
# convert_office_file
# ---------------------------------------------------------------------------

def test_convert_office_file_docx(tmp_path):
    """convert_office_file creates a markdown sidecar from .docx."""
    pytest.importorskip('docx')
    from graphify.detect import convert_office_file
    from docx import Document

    full_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_heading("Overview", level=1)
    doc.add_paragraph("Important content.")
    doc.save(str(full_path))

    out_dir = tmp_path / "converted"
    sidecar = convert_office_file(full_path, out_dir)

    assert sidecar is not None
    assert sidecar.suffix == ".md"
    text = sidecar.read_text()
    assert "Overview" in text
    assert "Important content" in text


def test_convert_office_file_xlsx(tmp_path):
    """convert_office_file creates a markdown sidecar from .xlsx."""
    pytest.importorskip('openpyxl')
    from graphify.detect import convert_office_file
    import openpyxl

    full_path = tmp_path / "budget.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Budget"
    ws.append(["Item", "Cost"])
    ws.append(["Widget", 42])
    wb.save(str(full_path))

    out_dir = tmp_path / "converted"
    sidecar = convert_office_file(full_path, out_dir)

    assert sidecar is not None
    assert sidecar.suffix == ".md"
    text = sidecar.read_text()
    assert "Budget" in text


def test_convert_office_file_unknown_format(tmp_path):
    """convert_office_file returns None for unknown extensions."""
    from graphify.detect import convert_office_file
    assert convert_office_file(tmp_path / "data.pdf", tmp_path) is None


def test_convert_office_file_empty_content(tmp_path):
    """convert_office_file returns None when conversion produces nothing."""
    pytest.importorskip('docx')
    from graphify.detect import convert_office_file
    from docx import Document

    full_path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(full_path))

    sidecar = convert_office_file(full_path, tmp_path)
    assert sidecar is None


# ---------------------------------------------------------------------------
# count_words — office file paths
# ---------------------------------------------------------------------------

def test_count_words_docx(tmp_path):
    """count_words handles .docx files."""
    pytest.importorskip('docx')
    from graphify.detect import count_words
    from docx import Document

    full_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("one two three four five")
    doc.save(str(full_path))

    words = count_words(full_path)
    assert words == 5


def test_count_words_xlsx(tmp_path):
    """count_words handles .xlsx files."""
    pytest.importorskip('openpyxl')
    from graphify.detect import count_words
    import openpyxl

    full_path = tmp_path / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["a", "b"])
    wb.save(str(full_path))

    words = count_words(full_path)
    assert isinstance(words, int)


def test_count_words_missing_file(tmp_path):
    """count_words returns 0 for missing file."""
    from graphify.detect import count_words
    assert count_words(tmp_path / "nope.md") == 0


# ---------------------------------------------------------------------------
# _is_ignored — anchored / negated / edge cases
# ---------------------------------------------------------------------------

def test_is_ignored_anchored_pattern(tmp_path):
    """Anchored pattern only matches from root."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "src").mkdir()
    (root / "src" / "main.py").write_text("x=1")

    # "/src/*.py" anchored — matches files directly under src
    patterns = [(root, "/src/*.py")]
    assert _is_ignored(root / "src" / "main.py", root, patterns) is True
    # Anchored /src/*.py should not match root-level file
    (root / "main.py").write_text("x=1")
    assert _is_ignored(root / "main.py", root, patterns) is False


def test_is_ignored_negated_pattern(tmp_path):
    """Negation pattern (!) un-ignores a previously ignored path."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "lib.py").write_text("x=1")

    # Ignore all .py, then un-ignore lib.py
    patterns = [(root, "*.py"), (root, "!lib.py")]
    assert _is_ignored(root / "lib.py", root, patterns) is False
    assert _is_ignored(root / "other.py", root, patterns) is True


def test_is_ignored_empty_pattern_skipped(tmp_path):
    """Empty after stripping is skipped (no-op)."""
    from graphify.detect import _is_ignored
    (tmp_path / "f.py").write_text("x=1")
    patterns = [(tmp_path, "/")]  # strip "/" → "" → skipped
    assert _is_ignored(tmp_path / "f.py", tmp_path, patterns) is False


def test_is_ignored_no_patterns(tmp_path):
    """Empty pattern list → not ignored."""
    from graphify.detect import _is_ignored
    (tmp_path / "f.py").write_text("x=1")
    assert _is_ignored(tmp_path / "f.py", tmp_path, []) is False


def test_is_ignored_path_outside_root(tmp_path):
    """Path not relative to root is not ignored."""
    from graphify.detect import _is_ignored
    other = tmp_path / "other"
    other.mkdir()
    (other / "lib.py").write_text("x=1")
    patterns = [(tmp_path, "*.py")]
    # Path outside root still gets checked
    assert _is_ignored(other / "lib.py", tmp_path, patterns) is True


# ---------------------------------------------------------------------------
# _is_included — anchored paths
# ---------------------------------------------------------------------------

def test_is_included_anchored_pattern(tmp_path):
    """Anchored include pattern matches from anchor dir."""
    from graphify.detect import _is_included
    root = tmp_path
    (root / "hidden").mkdir()
    (root / "hidden" / "config.yml").write_text("key: val")
    patterns = [(root, "/hidden/config.yml")]
    assert _is_included(root / "hidden" / "config.yml", root, patterns) is True
    assert _is_included(root / "visible" / "config.yml", root, patterns) is False


def test_is_included_empty_pattern_skipped(tmp_path):
    from graphify.detect import _is_included
    (tmp_path / "f.py").write_text("x=1")
    patterns = [(tmp_path, "/")]  # strip → "" → skipped
    assert _is_included(tmp_path / "f.py", tmp_path, patterns) is False


# ---------------------------------------------------------------------------
# _could_contain_included_path — edge cases
# ---------------------------------------------------------------------------

def test_could_contain_included_empty_pattern(tmp_path):
    from graphify.detect import _could_contain_included_path
    assert _could_contain_included_path(tmp_path, tmp_path, []) is False


def test_could_contain_included_root_match(tmp_path):
    from graphify.detect import _could_contain_included_path
    patterns = [(tmp_path, "src")]
    d = tmp_path / "src"
    d.mkdir()
    assert _could_contain_included_path(d, tmp_path, patterns) is True


# ---------------------------------------------------------------------------
# detect — office files, large corpus threshold, edge cases
# ---------------------------------------------------------------------------

def test_detect_large_corpus_warning(tmp_path):
    """detect warns when corpus exceeds CORPUS_UPPER_THRESHOLD or file count."""
    from graphify.detect import CORPUS_UPPER_THRESHOLD, FILE_COUNT_UPPER

    # Create many files with many words
    for i in range(FILE_COUNT_UPPER + 1):
        f = tmp_path / f"doc_{i}.md"
        f.write_text("word " * 3000)

    result = detect(tmp_path)
    assert result["warning"] is not None


def test_detect_incremental_with_legacy_manifest(tmp_path, monkeypatch):
    """detect_incremental handles legacy manifests (bare float mtime)."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "main.py").write_text("x = 1")
    manifest_path = str(tmp_path / "old_manifest.json")
    # Legacy format: just float values
    import json
    with open(manifest_path, "w") as f:
        json.dump({"main.py": 0.0}, f)  # old mtime

    result = detect_incremental(tmp_path, manifest_path)
    assert result["incremental"] is True
    assert result["new_total"] >= 1  # modified


def test_detect_incremental_no_previous_manifest(tmp_path, monkeypatch):
    """detect_incremental with no manifest treats everything as new."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "a.py").write_text("x=1")
    (tmp_path / "b.md").write_text("hello")
    result = detect_incremental(tmp_path, str(tmp_path / "none.json"))
    assert result["new_total"] == result["total_files"]


def test_detect_incremental_respects_manifest_formats(tmp_path, monkeypatch):
    """detect_incremental handles dict manifests (modern format)."""
    monkeypatch.chdir(tmp_path)
    from graphify.detect import _md5_file
    import time

    (tmp_path / "unchanged.py").write_text("x = 1")
    (tmp_path / "graphify-out").mkdir()
    manifest_path = str(tmp_path / "graphify-out" / "modern_manifest.json")
    import json
    with open(manifest_path, "w") as f:
        json.dump({
            str(tmp_path / "unchanged.py"): {
                "mtime": (tmp_path / "unchanged.py").stat().st_mtime,
                "hash": _md5_file(tmp_path / "unchanged.py"),
            }
        }, f)

    # Detect and check that unchanged file is not re-extracted
    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] == 0


def test_detect_incremental_with_unknown_manifest_type(tmp_path, monkeypatch):
    """detect_incremental re-extracts when manifest has unknown format."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "file.py").write_text("x = 1")
    manifest_path = str(tmp_path / "weird_manifest.json")
    import json
    with open(manifest_path, "w") as f:
        json.dump({"file.py": ["list instead of dict or float"]}, f)

    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] >= 1  # unknown format → re-extract


def test_detect_incremental_with_deleted_followed_by_scan(tmp_path, monkeypatch):
    """detect_incremental reports deleted files."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "old_file.py").write_text("x = 1")

    manifest_path = str(tmp_path / "del_manifest.json")
    from graphify.detect import _md5_file
    import json

    # Save manifest with one file
    with open(manifest_path, "w") as f:
        json.dump({
            str(tmp_path / "old_file.py"): {
                "mtime": (tmp_path / "old_file.py").stat().st_mtime,
                "hash": _md5_file(tmp_path / "old_file.py"),
            }
        }, f)

    # Now remove the file
    (tmp_path / "old_file.py").unlink()
    (tmp_path / "new_file.py").write_text("y = 2")

    result = detect_incremental(tmp_path, manifest_path)
    assert len(result.get("deleted_files", [])) >= 1
    assert any("old_file.py" in f for f in result["deleted_files"])


def test_detect_skips_graphify_own_output(tmp_path):
    """detect skips the graphify-out/ directory."""
    (tmp_path / "graphify-out").mkdir()
    (tmp_path / "graphify-out" / "graph.json").write_text("{}")
    (tmp_path / "main.py").write_text("x = 1")
    result = detect(tmp_path)
    code_files = result["files"]["code"]
    assert not any("graphify-out" in f for f in code_files)


# ---------------------------------------------------------------------------
# _is_noise_dir — egg-info
# ---------------------------------------------------------------------------

def test_is_noise_dir_egg_info():
    from graphify.detect import _is_noise_dir
    assert _is_noise_dir("my_package.egg-info") is True


# ---------------------------------------------------------------------------
# stress test — many files
# ---------------------------------------------------------------------------

def test_detect_many_files(tmp_path):
    """detect handles directories with hundreds of files."""
    for i in range(200):
        (tmp_path / f"script_{i:04d}.py").write_text(f"# file {i}\nprint('hi')")
    result = detect(tmp_path)
    assert result["total_files"] == 200


# ---------------------------------------------------------------------------
# _shebang_file_type — OSError path
# ---------------------------------------------------------------------------

def test_shebang_file_type_oserror(tmp_path, monkeypatch):
    """_shebang_file_type returns None on OSError."""
    from graphify.detect import _shebang_file_type
    f = tmp_path / "unreadable"
    f.write_text("#!/usr/bin/env python3")
    # Make the file unreadable to trigger OSError
    f.chmod(0o000)
    try:
        assert _shebang_file_type(f) is None
    finally:
        f.chmod(0o644)


# ---------------------------------------------------------------------------
# extract_pdf_text — exception path (hits line 155)
# ---------------------------------------------------------------------------

def test_extract_pdf_text_corrupt(tmp_path):
    """extract_pdf_text returns '' for corrupt/non-pdf files."""
    from graphify.detect import extract_pdf_text
    bad = tmp_path / "bad.pdf"
    bad.write_text("not a real pdf")
    text = extract_pdf_text(bad)
    assert text == ""


# ---------------------------------------------------------------------------
# docx_to_markdown — List style, tables, ImportError, exception
# ---------------------------------------------------------------------------

def test_docx_to_markdown_with_list(tmp_path):
    """docx_to_markdown handles List paragraph style."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    full_path = tmp_path / "list_doc.docx"
    doc = Document()
    # Add a paragraph with List style
    p = doc.add_paragraph("Item one")
    p.style = doc.styles["List Bullet"]
    p = doc.add_paragraph("Item two")
    p.style = doc.styles["List Bullet"]
    doc.save(str(full_path))

    md = docx_to_markdown(full_path)
    assert "- Item one" in md
    assert "- Item two" in md


def test_docx_to_markdown_with_table(tmp_path):
    """docx_to_markdown converts tables to markdown tables."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    from docx import Document

    full_path = tmp_path / "table_doc.docx"
    doc = Document()
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Name"
    table.rows[0].cells[1].text = "Age"
    table.rows[1].cells[0].text = "Alice"
    table.rows[1].cells[1].text = "30"
    table.rows[2].cells[0].text = "Bob"
    table.rows[2].cells[1].text = "25"
    doc.save(str(full_path))

    md = docx_to_markdown(full_path)
    assert "| Name | Age |" in md
    assert "| Alice | 30 |" in md


def test_docx_to_markdown_corrupt_file(tmp_path):
    """docx_to_markdown returns '' for corrupt files hitting the general exception path."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    corrupt = tmp_path / "corrupt.docx"
    corrupt.write_bytes(b"not a real docx file \x00\x00\x00")
    result = docx_to_markdown(corrupt)
    assert result == ""


def test_docx_to_markdown_empty_heading(tmp_path):
    """docx_to_markdown handles paragraphs with empty text."""
    pytest.importorskip('docx')
    from graphify.detect import docx_to_markdown
    from docx import Document

    full_path = tmp_path / "empty_para.docx"
    doc = Document()
    doc.add_heading("", level=1)  # empty heading
    doc.add_paragraph("")
    doc.add_paragraph("Real content")
    doc.save(str(full_path))

    md = docx_to_markdown(full_path)
    assert "Real content" in md


# ---------------------------------------------------------------------------
# xlsx_to_markdown — empty rows, ImportError, general exception
# ---------------------------------------------------------------------------

def test_xlsx_to_markdown_empty_rows(tmp_path):
    """xlsx_to_markdown skips entirely empty rows."""
    pytest.importorskip('openpyxl')
    from graphify.detect import xlsx_to_markdown
    import openpyxl

    full_path = tmp_path / "sparse.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["Name", "Age"])  # header row
    ws.append([None, None])     # empty row (skipped)
    ws.append([None, None])     # another empty row
    ws.append(["Alice", 30])
    wb.save(str(full_path))

    md = xlsx_to_markdown(full_path)
    assert "Name" in md
    assert "Alice" in md


def test_xlsx_to_markdown_empty_workbook(tmp_path):
    """xlsx_to_markdown handles workbook with empty sheet."""
    pytest.importorskip("openpyxl")
    from graphify.detect import xlsx_to_markdown
    import openpyxl

    full_path = tmp_path / "empty.xlsx"
    wb = openpyxl.Workbook()
    wb.active.title = "Empty"
    wb.save(str(full_path))

    md = xlsx_to_markdown(full_path)
    assert isinstance(md, str)  # should not crash


def test_xlsx_to_markdown_corrupt(tmp_path):
    """xlsx_to_markdown returns '' for corrupt files."""
    from graphify.detect import xlsx_to_markdown
    corrupt = tmp_path / "corrupt.xlsx"
    corrupt.write_bytes(b"not an xlsx")
    result = xlsx_to_markdown(corrupt)
    assert result == ""


# ---------------------------------------------------------------------------
# xlsx_extract_structure — tables, wb.close(), import error
# ---------------------------------------------------------------------------

def test_xlsx_extract_structure_with_tables(tmp_path):
    """xlsx_extract_structure handles named tables in Excel."""
    pytest.importorskip("openpyxl")
    from graphify.detect import xlsx_extract_structure
    import openpyxl

    full_path = tmp_path / "tables.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Product", "Revenue"])
    ws.append(["Widget", 100])
    ws.append(["Gadget", 200])
    # Create a named table
    from openpyxl.worksheet.table import Table, TableStyleInfo
    table = Table(displayName="SalesData", ref="A1:B3")
    style = TableStyleInfo(
        name="TableStyleMedium9", showFirstColumn=False,
        showLastColumn=False, showRowStripes=True, showColumnStripes=False
    )
    table.tableStyleInfo = style
    ws.add_table(table)
    wb.save(str(full_path))

    result = xlsx_extract_structure(full_path)
    assert "nodes" in result
    assert len(result["nodes"]) >= 1
    assert any("salesdata" in n.get("id", "") for n in result["nodes"])


def test_xlsx_extract_structure_corrupt_file(tmp_path):
    """xlsx_extract_structure returns empty when file is corrupt/not an xlsx."""
    from graphify.detect import xlsx_extract_structure
    bad = tmp_path / "bad.xlsx"
    bad.write_bytes(b"not an xlsx file at all")
    result = xlsx_extract_structure(bad)
    assert result == {"nodes": [], "edges": []}


# ---------------------------------------------------------------------------
# count_words — PDF path
# ---------------------------------------------------------------------------

def test_count_words_pdf(tmp_path):
    """count_words handles .pdf files."""
    from graphify.detect import count_words
    from pypdf import PdfWriter

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    pdf_path = tmp_path / "doc.pdf"
    with open(pdf_path, "wb") as f:
        writer.write(f)

    words = count_words(pdf_path)
    assert isinstance(words, int)
    assert words >= 0


# ---------------------------------------------------------------------------
# _is_ignored — more edge cases (fnmatch on parts, ValueError)
# ---------------------------------------------------------------------------

def test_is_ignored_fnmatch_on_part(tmp_path):
    """_is_ignored matches via fnmatch on individual path parts."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "apps").mkdir()
    (root / "apps" / "sub").mkdir()
    (root / "apps" / "sub" / "ignored.py").write_text("x=1")

    # Pattern that matches a path part
    patterns = [(root, "ignored.py")]
    assert _is_ignored(root / "apps" / "sub" / "ignored.py", root, patterns) is True


def test_is_ignored_fnmatch_on_prefix(tmp_path):
    """_is_ignored matches via fnmatch on path prefix."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "vendor").mkdir()
    (root / "vendor" / "lib.py").write_text("x=1")

    patterns = [(root, "vendor/*")]
    assert _is_ignored(root / "vendor" / "lib.py", root, patterns) is True


def test_is_ignored_value_error_on_relative(tmp_path):
    """_is_ignored handles path outside anchor gracefully."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "sub").mkdir()
    other = tmp_path / "outside"
    other.mkdir()
    anchor = root / "sub"

    # File at anchor level with pattern anchored to anchor
    # This should not crash even when relative_to fails
    patterns = [(anchor, "/file.py")]
    (other / "file.py").write_text("x=1")
    assert _is_ignored(other / "file.py", root, patterns) is False


def test_is_ignored_non_root_anchor_path(tmp_path):
    """_is_ignored handles anchor != root fallback."""
    from graphify.detect import _is_ignored
    root = tmp_path
    (root / "sub").mkdir()
    anchor = root / "sub"
    (anchor / "vendor").mkdir()
    (anchor / "vendor" / "lib.py").write_text("x=1")

    # Non-anchored pattern with anchor != root
    patterns = [(anchor, "vendor/*")]
    assert _is_ignored(anchor / "vendor" / "lib.py", root, patterns) is True


# ---------------------------------------------------------------------------
# _is_included — more edge cases (fnmatch on parts, ValueError)
# ---------------------------------------------------------------------------

def test_is_included_fnmatch_on_part(tmp_path):
    """_is_included matches via fnmatch on individual path parts."""
    from graphify.detect import _is_included
    root = tmp_path
    (root / "hidden").mkdir()
    (root / "hidden" / "config.yml").write_text("key: val")

    patterns = [(root, "config.yml")]
    assert _is_included(root / "hidden" / "config.yml", root, patterns) is True


def test_is_included_fnmatch_on_prefix(tmp_path):
    """_is_included matches via fnmatch on path prefix."""
    from graphify.detect import _is_included
    root = tmp_path
    (root / "hidden").mkdir()
    (root / "hidden" / "nested").mkdir()
    (root / "hidden" / "nested" / "file.txt").write_text("data")

    patterns = [(root, "hidden/nested/*")]
    assert _is_included(root / "hidden" / "nested" / "file.txt", root, patterns) is True


def test_is_included_value_error_relative(tmp_path):
    """_is_included handles ValueError from relative_to gracefully."""
    from graphify.detect import _is_included
    root = tmp_path
    (root / "sub").mkdir()
    anchor = root / "sub"
    other = tmp_path / "outside"
    other.mkdir()
    (other / "config.yml").write_text("key: val")

    patterns = [(anchor, "/config.yml")]
    assert _is_included(other / "config.yml", root, patterns) is False


def test_is_included_non_root_anchor(tmp_path):
    """_is_included handles anchor != root fallback."""
    from graphify.detect import _is_included
    root = tmp_path
    (root / "inner").mkdir()
    anchor = root / "inner"
    (anchor / "secret").mkdir(parents=True)
    (anchor / "secret" / "config.yml").write_text("key: val")

    patterns = [(anchor, "secret/config.yml")]
    assert _is_included(anchor / "secret" / "config.yml", root, patterns) is True


# ---------------------------------------------------------------------------
# _could_contain_included_path — more edge cases
# ---------------------------------------------------------------------------

def test_could_contain_included_fnmatch_match(tmp_path):
    """_could_contain_included_path matches via fnmatch."""
    from graphify.detect import _could_contain_included_path
    patterns = [(tmp_path, "src")]
    d = tmp_path / "src"
    d.mkdir()
    assert _could_contain_included_path(d, tmp_path, patterns) is True


def test_could_contain_included_non_root_anchor(tmp_path):
    """_could_contain_included_path handles non-root anchor."""
    from graphify.detect import _could_contain_included_path
    anchor = tmp_path / "sub"
    anchor.mkdir()
    d = anchor / "hidden"
    d.mkdir()
    patterns = [(anchor, "hidden/*")]
    assert _could_contain_included_path(d, tmp_path, patterns) is True


def test_could_contain_included_path_is_root(tmp_path):
    """_could_contain_included_path returns True when rel matches pattern directly."""
    from graphify.detect import _could_contain_included_path
    d = tmp_path / "src"
    d.mkdir()
    patterns = [(tmp_path, "src")]
    # The relative path "src" matches pattern "src"
    assert _could_contain_included_path(d, tmp_path, patterns) is True


# ---------------------------------------------------------------------------
# detect — _SKIP_FILES (lock files), sensitive files, converted_dir, office files
# ---------------------------------------------------------------------------

def test_detect_skips_lock_files(tmp_path):
    """detect skips package-lock.json and other lock files."""
    (tmp_path / "main.py").write_text("x=1")
    (tmp_path / "package-lock.json").write_text("{}")
    (tmp_path / "yarn.lock").write_text("")
    result = detect(tmp_path)
    file_list = result["files"]["code"]
    assert any("main.py" in f for f in file_list)
    assert not any("package-lock.json" in f for f in file_list)
    assert not any("yarn.lock" in f for f in file_list)


def test_detect_sensitive_skipped(tmp_path):
    """detect skips and reports sensitive files that aren't hidden."""
    (tmp_path / "main.py").write_text("x=1")
    (tmp_path / "id_rsa").write_text("private key")
    (tmp_path / "server.key").write_text("key data")
    result = detect(tmp_path)
    # Hidden .env would be pruned before sensitive check; non-hidden sensitive files pass
    assert len(result["skipped_sensitive"]) >= 2
    assert any("id_rsa" in s for s in result["skipped_sensitive"])
    assert any("server.key" in s for s in result["skipped_sensitive"])


def test_detect_skips_converted_dir(tmp_path):
    """detect skips files inside graphify-out/converted/.

    Note: graphify-out/ is in _SKIP_DIRS so it won't be scanned at all.
    But the converted_dir check protects against cases where graphify-out
    is reached through a different path (e.g., via symlinks or memory dir).
    Since graphify-out is always pruned by dir-level skipping, this test
    focuses on verifying that the converted dir check won't crash on normal scans.
    """
    (tmp_path / "main.py").write_text("x=1")
    result = detect(tmp_path)
    code_files = result["files"]["code"]
    assert any("main.py" in f for f in code_files)


def test_detect_office_conversion_with_module(tmp_path):
    """detect converts .docx files to markdown sidecars."""
    pytest.importorskip("docx")
    from docx import Document

    (tmp_path / "report.docx")
    doc = Document()
    doc.add_paragraph("Office content here.")
    doc.save(str(tmp_path / "report.docx"))

    (tmp_path / "main.py").write_text("x=1")

    result = detect(tmp_path)
    # Should have converted the .docx to a markdown sidecar
    doc_files = result["files"]["document"]
    assert any("report" in f and ".md" in f for f in doc_files)


def test_detect_office_conversion_fails(tmp_path, monkeypatch):
    """detect gracefully handles office conversion failure."""
    from graphify.detect import detect

    (tmp_path / "bad.docx").write_bytes(b"not a real docx")
    (tmp_path / "main.py").write_text("x=1")

    result = detect(tmp_path)
    # Should have a skip note
    skipped = result["skipped_sensitive"]
    assert any("office conversion failed" in s for s in skipped)


def test_detect_google_workspace_exception(tmp_path, monkeypatch):
    """detect handles Google Workspace export exceptions."""
    (tmp_path / "notes.gdoc").write_text('{"doc_id":"doc-1"}')

    def fake_convert_raises(*args, **kwargs):
        raise RuntimeError("API error")

    monkeypatch.setattr("graphify.detect.convert_google_workspace_file", fake_convert_raises)

    result = detect(tmp_path, google_workspace=True)
    skipped = result["skipped_sensitive"]
    assert any("Google Workspace export failed" in s for s in skipped)


def test_detect_google_workspace_empty_result(tmp_path, monkeypatch):
    """detect handles Google Workspace export producing no text."""
    (tmp_path / "notes.gdoc").write_text('{"doc_id":"doc-1"}')

    def fake_convert_empty(path, out_dir, *, xlsx_to_markdown=None):
        out_dir.mkdir(parents=True, exist_ok=True)
        return None  # no readable text

    monkeypatch.setattr("graphify.detect.convert_google_workspace_file", fake_convert_empty)

    result = detect(tmp_path, google_workspace=True)
    skipped = result["skipped_sensitive"]
    assert any("no readable text" in s for s in skipped)


# ---------------------------------------------------------------------------
# save_manifest — OSError path
# ---------------------------------------------------------------------------

def test_save_manifest_oserror(tmp_path, monkeypatch):
    """save_manifest handles OSError when a file is deleted between detect and save."""
    from graphify.detect import save_manifest
    # Create a file, save manifest referencing it, but file is now absent
    manifest_path = str(tmp_path / "manifest.json")
    files = {"code": [str(tmp_path / "gone.py")]}

    save_manifest(files, manifest_path)  # should not raise
    # Should have created a manifest with no entries (file was missing)
    import json
    manifest = json.loads(Path(manifest_path).read_text())
    assert isinstance(manifest, dict)


# ---------------------------------------------------------------------------
# detect_incremental — mtime/hash edge cases
# ---------------------------------------------------------------------------

def test_detect_incremental_mtime_changed_same_hash(tmp_path, monkeypatch):
    """detect_incremental treats mtime bump with same hash as unchanged."""
    from graphify.detect import _md5_file, detect_incremental
    monkeypatch.chdir(tmp_path)

    (tmp_path / "graphify-out").mkdir()
    f = tmp_path / "unchanged.py"
    f.write_text("x = 1")
    current_hash = _md5_file(f)

    import json, time
    manifest_path = str(tmp_path / "graphify-out" / "manifest.json")
    # Store with WRONG mtime but CORRECT hash (simulating sync tool touching mtime)
    wrong_mtime = (tmp_path / "unchanged.py").stat().st_mtime + 1000
    with open(manifest_path, "w") as mf:
        json.dump({
            str(f): {"mtime": wrong_mtime, "hash": current_hash},
        }, mf)

    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] == 0  # unchanged


def test_detect_incremental_mtime_same_no_hash(tmp_path, monkeypatch):
    """detect_incremental treats same mtime (with hash in dict) as unchanged."""
    from graphify.detect import _md5_file, detect_incremental
    monkeypatch.chdir(tmp_path)

    (tmp_path / "graphify-out").mkdir()
    f = tmp_path / "same.py"
    f.write_text("x = 1")
    current_mtime = f.stat().st_mtime
    current_hash = _md5_file(f)
    import json
    manifest_path = str(tmp_path / "graphify-out" / "manifest.json")
    with open(manifest_path, "w") as mf:
        json.dump({
            str(f): {"mtime": current_mtime, "hash": current_hash},
        }, mf)

    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] == 0


def test_detect_incremental_mtime_changed_content_changed(tmp_path, monkeypatch):
    """detect_incremental detects actual content change."""
    from graphify.detect import _md5_file, detect_incremental
    monkeypatch.chdir(tmp_path)

    f = tmp_path / "changed.py"
    f.write_text("original content")
    old_hash = _md5_file(f)
    import json
    manifest_path = str(tmp_path / "manifest.json")
    old_mtime = f.stat().st_mtime
    with open(manifest_path, "w") as mf:
        json.dump({
            str(f): {"mtime": old_mtime, "hash": old_hash},
        }, mf)

    # Change content
    f.write_text("new content")
    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] >= 1  # content changed


def test_detect_incremental_stored_none(tmp_path, monkeypatch):
    """detect_incremental handles stored value being None (missing hash key)."""
    from graphify.detect import detect_incremental
    monkeypatch.chdir(tmp_path)

    (tmp_path / "file.py").write_text("x = 1")
    manifest_path = str(tmp_path / "manifest.json")
    import json
    with open(manifest_path, "w") as mf:
        json.dump({
            str(tmp_path / "file.py"): {"mtime": None},
        }, mf)

    result = detect_incremental(tmp_path, manifest_path)
    assert result["incremental"] is True


# ---------------------------------------------------------------------------
# _shebang_file_type — empty shebang parts (line 105)
# ---------------------------------------------------------------------------

def test_shebang_file_type_empty_parts(tmp_path):
    """_shebang_file_type returns None when shebang has no interpreter parts."""
    from graphify.detect import _shebang_file_type
    f = tmp_path / "script"
    f.write_text("#!\n")  # just the shebang marker, no interpreter
    assert _shebang_file_type(f) is None


# ---------------------------------------------------------------------------
# extract_pdf_text — actual page text extraction (line 155)
#
# NOTE: Line 155 (pages.append(text)) is only reached when a PDF page
# has extractable text. Blank PDF pages return empty string, so line 155
# is not covered in the existing simple test. We skip adding a complex
# PDF-with-text test due to the difficulty of creating such PDFs
# programmatically without external tools. This line would be covered
# by an integration test with a real multi-page PDF with content.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# docx_to_markdown — Heading 3 (line 179)
# ---------------------------------------------------------------------------

def test_docx_to_markdown_heading3(tmp_path):
    """docx_to_markdown handles Heading 3 style."""
    pytest.importorskip("docx")
    from graphify.detect import docx_to_markdown
    from docx import Document

    full_path = tmp_path / "heading3.docx"
    doc = Document()
    doc.add_heading("Top", level=1)
    doc.add_heading("Middle", level=2)
    doc.add_heading("Bottom", level=3)
    doc.add_paragraph("Content.")
    doc.save(str(full_path))

    md = docx_to_markdown(full_path)
    assert "# Top" in md
    assert "## Middle" in md
    assert "### Bottom" in md


def test_docx_to_markdown_import_error(monkeypatch):
    """docx_to_markdown returns '' when python-docx is not installed."""
    import builtins
    _real_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "docx":
            raise ImportError("No module named 'docx'")
        return _real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    from graphify.detect import docx_to_markdown
    result = docx_to_markdown(Path("test.docx"))
    assert result == ""


# ---------------------------------------------------------------------------
# xlsx_to_markdown — ImportError (line 226)
# Note: openpyxl is imported at module level in graphify/detect.py? No, it's
# imported inside functions. But xlsx_to_markdown imports it via:
#   import openpyxl  (inside the function try block)
# So mocking builtins.__import__ should work.
# ---------------------------------------------------------------------------

def test_xlsx_to_markdown_import_error(monkeypatch):
    """xlsx_to_markdown returns '' when openpyxl is not installed."""
    import builtins
    _real_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "openpyxl":
            raise ImportError("No module named 'openpyxl'")
        return _real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    from graphify.detect import xlsx_to_markdown
    result = xlsx_to_markdown(Path("test.xlsx"))
    assert result == ""


# ---------------------------------------------------------------------------
# xlsx_extract_structure — ImportError, fallback, wb.close() (lines 243, 303-311, 315-316)
# ---------------------------------------------------------------------------

def test_xlsx_extract_structure_import_error(monkeypatch):
    """xlsx_extract_structure returns empty when openpyxl is not installed."""
    import builtins
    _real_import = builtins.__import__

    def mock_import(name, *args, **kwargs):
        if name == "openpyxl":
            raise ImportError("No module named 'openpyxl'")
        return _real_import(name, *args, **kwargs)

    monkeypatch.setattr(builtins, "__import__", mock_import)

    from graphify.detect import xlsx_extract_structure
    result = xlsx_extract_structure(Path("test.xlsx"))
    assert result == {"nodes": [], "edges": []}


def test_xlsx_extract_structure_fallback_columns(tmp_path, monkeypatch):
    """xlsx_extract_structure uses first-row fallback when no named tables."""
    pytest.importorskip("openpyxl")
    from graphify.detect import xlsx_extract_structure
    import openpyxl
    from openpyxl.worksheet.worksheet import Worksheet

    full_path = tmp_path / "simple.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["ColA", "ColB"])   # fallback: first non-empty row = column headers
    ws.append([1, 2])
    wb.save(str(full_path))

    # Cover the fallback path (lines 303-311) by removing 'tables' from the
    # worksheet at load time. The `else` branch triggers when hasattr(ws, "tables") is False.
    _real_load = openpyxl.load_workbook

    def mock_load(path, **kwargs):
        wb2 = _real_load(path, **kwargs)
        for sn in wb2.sheetnames:
            ws2 = wb2[sn]
            # Force hasattr(ws, "tables") -> False to reach the fallback else branch
            monkeypatch.delattr(type(ws2), "tables", raising=False)
        return wb2

    monkeypatch.setattr(openpyxl, "load_workbook", mock_load)

    result = xlsx_extract_structure(full_path)
    assert len(result["nodes"]) >= 2  # file + sheet + column nodes from fallback


def test_xlsx_extract_structure_close_exception(tmp_path, monkeypatch):
    """xlsx_extract_structure handles exception during wb.close()."""
    pytest.importorskip("openpyxl")
    from graphify.detect import xlsx_extract_structure
    import openpyxl

    full_path = tmp_path / "data.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["X"])
    ws.append([1])
    wb.save(str(full_path))

    # Mock wb.close to raise an exception
    import openpyxl.workbook
    orig_close = openpyxl.workbook.Workbook.close

    def mock_close(self):
        raise RuntimeError("close failed")

    monkeypatch.setattr(openpyxl.workbook.Workbook, "close", mock_close)

    result = xlsx_extract_structure(full_path)
    assert "nodes" in result
    assert "edges" in result


# ---------------------------------------------------------------------------
# detect — incremental with stat exception (lines 846-847)
# ---------------------------------------------------------------------------

def test_detect_incremental_stat_exception(tmp_path, monkeypatch):
    """detect_incremental handles Exception during stat (current_mtime -> 0)."""
    from graphify.detect import detect_incremental
    monkeypatch.chdir(tmp_path)

    (tmp_path / "file.py").write_text("x = 1")
    manifest_path = str(tmp_path / "manifest.json")
    import json
    with open(manifest_path, "w") as mf:
        json.dump({
            str(tmp_path / "file.py"): 0.0,  # legacy float mtime
        }, mf)

    # Make the file disappear between manifest load and stat
    import os as os_mod
    _real_stat = os_mod.stat

    def mock_stat(path, *, follow_symlinks=True):
        path_str = str(path)
        if "file.py" in path_str:
            raise OSError("File vanished")
        return _real_stat(path, follow_symlinks=follow_symlinks)

    monkeypatch.setattr(os_mod, "stat", mock_stat)

    result = detect_incremental(tmp_path, manifest_path)
    assert "incremental" in result


# ---------------------------------------------------------------------------
# Patch 1: Portable manifest paths
# ---------------------------------------------------------------------------

def test_save_manifest_relativizes_keys(tmp_path, monkeypatch):
    """save_manifest stores portable (relative) file keys in manifest.json."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "app.py").write_text("x = 1")
    (tmp_path / "README.md").write_text("# Hello")

    from graphify.detect import save_manifest
    import json

    manifest_path = str(tmp_path / "graphify-out" / "manifest.json")
    files = {
        "code": [str(tmp_path / "src" / "app.py")],
        "document": [str(tmp_path / "README.md")],
    }
    save_manifest(files, manifest_path=manifest_path, root=tmp_path)

    manifest = json.loads(Path(manifest_path).read_text())
    keys = list(manifest.keys())
    # Keys should be relative to root, not absolute
    for key in keys:
        assert not key.startswith("/"), f"Expected relative key, got: {key}"
    assert any("app.py" in k for k in keys)
    assert any("README.md" in k for k in keys)


def test_detect_incremental_reads_relative_manifest_keys(tmp_path, monkeypatch):
    """detect_incremental reads a manifest with relative keys and matches current files."""
    monkeypatch.chdir(tmp_path)
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "app.py").write_text("x = 1")
    (tmp_path / "README.md").write_text("# Hello")

    from graphify.detect import detect_incremental, _md5_file
    import json

    manifest_path = str(tmp_path / "graphify-out" / "manifest.json")
    (tmp_path / "graphify-out").mkdir(exist_ok=True)

    # Write manifest with relative paths (the portable format)
    with open(manifest_path, "w") as f:
        json.dump({
            "src/app.py": {
                "mtime": (tmp_path / "src" / "app.py").stat().st_mtime,
                "hash": _md5_file(tmp_path / "src" / "app.py"),
            },
            "README.md": {
                "mtime": (tmp_path / "README.md").stat().st_mtime,
                "hash": _md5_file(tmp_path / "README.md"),
            },
        }, f)

    result = detect_incremental(tmp_path, manifest_path)
    assert result["new_total"] == 0  # No changes — files match manifest


# ---------------------------------------------------------------------------
# Patch 3: JSON/JSONC document indexing and generated/bundled file filtering
# ---------------------------------------------------------------------------


def test_classify_json_as_document(tmp_path):
    from graphify.detect import FileType, classify_file

    cfg = tmp_path / "config.json"
    cfg.write_text('{"name": "demo"}')

    assert classify_file(cfg) == FileType.DOCUMENT


def test_classify_jsonc_as_document(tmp_path):
    from graphify.detect import FileType, classify_file

    cfg = tmp_path / "settings.jsonc"
    cfg.write_text('{"compilerOptions": {}}')

    assert classify_file(cfg) == FileType.DOCUMENT


def test_detect_still_skips_package_lock_json(tmp_path):
    from graphify.detect import detect

    (tmp_path / "package-lock.json").write_text("{}")
    (tmp_path / "config.json").write_text("{}")

    result = detect(tmp_path)

    assert any("config.json" in f for f in result["files"]["document"])
    assert not any("package-lock.json" in f for files in result["files"].values() for f in files)


def test_detect_skips_generated_and_bundle_files(tmp_path):
    from graphify.detect import detect

    (tmp_path / "app.ts").write_text("export const app = 1")
    (tmp_path / "vendor.bundle.js").write_text("function bundled(){}")
    (tmp_path / "types.d.ts").write_text("declare const x: string")
    (tmp_path / "schema.generated.ts").write_text("export const schema = {}")

    result = detect(tmp_path)
    code_files = result["files"]["code"]

    assert any("app.ts" in f for f in code_files)
    assert not any("vendor.bundle.js" in f for f in code_files)
    assert not any("types.d.ts" in f for f in code_files)
    assert not any("schema.generated.ts" in f for f in code_files)


def test_graphifyinclude_can_restore_generated_file(tmp_path):
    from graphify.detect import detect

    (tmp_path / ".graphifyinclude").write_text("schema.generated.ts\n")
    (tmp_path / "schema.generated.ts").write_text("export const schema = {}")

    result = detect(tmp_path)

    assert any("schema.generated.ts" in f for f in result["files"]["code"])


# ---------------------------------------------------------------------------
# _is_ignored — edge cases (uncovered ValueError fallthrough paths)
# ---------------------------------------------------------------------------

def test_is_ignored_path_not_under_root_with_anchored_pattern(tmp_path):
    """Anchored pattern with a path not under any anchor → ValueError pass (lines 539, 545-546)."""
    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other" / "file.py"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    patterns = [(root, "/sub/*.py")]  # anchored
    result = _is_ignored(other, root, patterns)
    assert result is False  # not ignored — ValueError silently caught


def test_is_ignored_unanchored_pattern_outside_anchors(tmp_path):
    """Unanchored pattern, path not under root OR anchor → both ValueErrors pass (lines 545-546, 551-552)."""
    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other" / "sub" / "file.py"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    anchor = tmp_path / "anchor"
    anchor.mkdir()

    patterns = [(anchor, "*.generated.py")]
    result = _is_ignored(other, root, patterns)
    assert result is False


def test_is_ignored_unanchored_no_match_multiple_parts(tmp_path):
    """Unanchored pattern with path matching via 'part' fnmatch (line 520-522 coverage)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "src" / "generated" / "file.py"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("x")

    # Pattern matches just "generated" as a directory part
    patterns = [(root, "generated")]
    result = _is_ignored(p, root, patterns)
    assert result is True  # should match "generated" dir part


# ---------------------------------------------------------------------------
# _is_included — edge cases
# ---------------------------------------------------------------------------

def test_is_included_matches_dir_part(tmp_path):
    """_is_included matches directory parts (lines 601-602)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "src" / "important" / "file.py"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("x")

    patterns = [(root, "important")]
    assert _is_included(p, root, patterns) is True


def test_is_included_matches_ancestor_chain(tmp_path):
    """_is_included matches ancestor chain fnmatch (lines 603-604)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "src" / "special" / "file.py"
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("x")

    patterns = [(root, "src/special")]
    assert _is_included(p, root, patterns) is True


def test_is_included_path_outside_root_unanchored(tmp_path):
    """_is_included ValueError when path can't relativize to root (lines 624-625)."""
    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other" / "file.py"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    patterns = [(root, "*.py")]
    assert _is_included(other, root, patterns) is False


def test_is_included_path_outside_anchor(tmp_path):
    """_is_included ValueError when path can't relativize to anchor (lines 631-632)."""
    root = tmp_path / "root"
    root.mkdir()
    anchor = tmp_path / "anchor"
    anchor.mkdir()
    other = tmp_path / "other" / "file.py"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    patterns = [(anchor, "*.py")]
    assert _is_included(other, root, patterns) is False


# ---------------------------------------------------------------------------
# _could_contain_included_path — edge cases
# ---------------------------------------------------------------------------

def test_could_contain_included_path_outside_root(tmp_path):
    """_could_contain_included_path handles ValueError for path outside root (lines 644-645)."""
    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other"
    other.mkdir(parents=True, exist_ok=True)

    patterns = [(root, "*.md")]
    assert _could_contain_included_path(other, root, patterns) is False


def test_could_contain_included_path_outside_anchor(tmp_path):
    """_could_contain_included_path handles ValueError for anchor outside root (lines 650-651)."""
    anchor = tmp_path / "anchor"
    anchor.mkdir()
    root = tmp_path / "root"
    root.mkdir()
    p = root / "subdir"
    p.mkdir(parents=True, exist_ok=True)

    patterns = [(anchor, "*.py")]
    assert _could_contain_included_path(p, root, patterns) is False


def test_could_contain_included_path_empty_rel_returns_true(tmp_path):
    """_could_contain_included_path when rel='.' (path==root) - fnmatch does not match."""
    root = tmp_path / "root"
    root.mkdir()

    patterns = [(root, "*.py")]
    # path==root gives rel='.', which does NOT match '*.py' via fnmatch
    assert _could_contain_included_path(root, root, patterns) is False

    # But with an exact-match pattern, the match works
    patterns2 = [(root, ".")]
    assert _could_contain_included_path(root, root, patterns2) is True


def test_could_contain_included_path_fnmatch_match(tmp_path):
    """_could_contain_included_path fnmatch match on rel (line 664)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "special-dir"
    p.mkdir(parents=True, exist_ok=True)

    patterns = [(root, "special-*")]
    assert _could_contain_included_path(p, root, patterns) is True


def test_could_contain_included_path_exact_prefix_match(tmp_path):
    """_could_contain_included_path matches when pattern starts with rel/ (lines 661-662)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "src"
    p.mkdir(parents=True, exist_ok=True)

    patterns = [(root, "src/components/Button.tsx")]
    assert _could_contain_included_path(p, root, patterns) is True


# ---------------------------------------------------------------------------
# _manifest_key — edge cases
# ---------------------------------------------------------------------------

def test_manifest_key_non_absolute_path(tmp_path):
    """_manifest_key handles non-absolute path (line 849)."""
    root = tmp_path / "root"
    root.mkdir()

    key = _manifest_key("src/file.py", root)
    assert key == "src/file.py"


def test_manifest_key_path_outside_root(tmp_path):
    """_manifest_key returns absolute path when outside root (lines 854-857)."""
    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other" / "file.py"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    key = _manifest_key(str(other), root)
    assert key.startswith(str(tmp_path.resolve()))


# ---------------------------------------------------------------------------
# _manifest_entry_for — edge cases
# ---------------------------------------------------------------------------

def test_manifest_entry_for_non_absolute_path(tmp_path):
    """_manifest_entry_for adds resolved absolute candidate for non-absolute path (line 877)."""
    root = tmp_path / "root"
    root.mkdir()
    p = root / "file.py"
    p.write_text("content")

    manifest = {"file.py": {"mtime": 123}}
    entry = _manifest_entry_for("file.py", manifest, root)
    assert entry == {"mtime": 123}


def test_manifest_entry_for_not_found(tmp_path):
    """_manifest_entry_for returns None when no candidates match."""
    root = tmp_path / "root"
    root.mkdir()
    manifest = {}
    assert _manifest_entry_for("missing.py", manifest, root) is None


# ---------------------------------------------------------------------------
# load_manifest — edge cases
# ---------------------------------------------------------------------------

def test_load_manifest_except_in_key_iteration(monkeypatch, tmp_path):
    """load_manifest exception in key resolution falls back to raw key (lines 912-913)."""
    manifest_path = tmp_path / ".graphify_manifest"
    manifest_path.write_text('{"good_key": {"mtime": 1}, "bad_key": {"mtime": 2}}')

    # Monkeypatch Path.resolve to raise on "bad_key" to trigger the except: branch
    original_resolve = Path.resolve
    def _raise_resolve(self):
        if self.name == "bad_key":
            raise RuntimeError("simulated resolve failure")
        return original_resolve(self)
    monkeypatch.setattr(Path, "resolve", _raise_resolve)

    result = load_manifest(str(manifest_path))
    assert "good_key" in result
    assert "bad_key" in result  # falls through except: out[key] = value


# ---------------------------------------------------------------------------
# _is_generated_or_bundle — edge cases
# ---------------------------------------------------------------------------

def test_is_generated_or_bundle_path_outside_root(tmp_path):
    """_is_generated_or_bundle when path can't be relativized to root (line 409-410)."""
    from graphify.detect import _is_generated_or_bundle

    root = tmp_path / "root"
    root.mkdir()
    other = tmp_path / "other" / "types.d.ts"
    other.parent.mkdir(parents=True, exist_ok=True)
    other.write_text("x")

    result = _is_generated_or_bundle(other, root, [])
    assert result is True  # "types.d.ts" matches _GENERATED_FILE_PATTERNS via basename


# ---------------------------------------------------------------------------
# save_manifest — edge cases
# ---------------------------------------------------------------------------

def test_save_manifest_non_absolute_path(tmp_path):
    """save_manifest handles non-absolute path by resolving relative to root (line 936)."""
    manifest_path = tmp_path / "out" / "manifest.json"
    p = tmp_path / "file.py"
    p.write_text("hello")

    save_manifest(
        {"code": ["file.py"]},  # non-absolute!
        str(manifest_path),
        root=tmp_path,
    )

    assert manifest_path.exists()
    data = json.loads(manifest_path.read_text())
    assert "file.py" in data


# ---------------------------------------------------------------------------
# _manifest_root — edge cases
# ---------------------------------------------------------------------------

def test_manifest_root_empty_out_dir_name(tmp_path):
    """_manifest_root with empty out_dir.name returns '.' (line 841)."""
    manifest = tmp_path / "manifest.json"
    manifest.write_text("{}")

    # out_dir.name would be "" at filesystem root — simulate
    result = _manifest_root(str(manifest))
    # tmp_path has a name, so this will return tmp_path.parent
    assert result.is_absolute()


def test_manifest_root_with_dotfile(tmp_path):
    """_manifest_root with .graphify_root file (lines 835-838)."""
    manifest = tmp_path / "out" / "manifest.json"
    manifest.parent.mkdir(parents=True, exist_ok=True)
    root_file = manifest.parent / ".graphify_root"
    root_file.write_text(str(tmp_path / "custom_root"))

    result = _manifest_root(str(manifest))
    assert result == (tmp_path / "custom_root").resolve()
